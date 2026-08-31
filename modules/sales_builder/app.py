"""
Smart 1 Sales Builder — the Proposal Builder
============================================
Dashboard, proposal wizard and IO hand-off, in one module.

**This is now the only proposal builder.** There were two: this one, and a
standalone block-based builder at /sales/proposals. They shared no code, no
storage and no idea of what a campaign is, so the same client could be quoted
differently depending on which tool a rep happened to open, and only this one
produced anything an insertion order could read. /sales/proposals now redirects
here; its saved proposals stay readable and can be imported as quotes (see
`/api/legacy/proposals`).

What came over from the standalone builder:

- the industry library (channels, demand triggers, intro copy) — now
  `hub.industries`, used for the default proposal sections and the AI draft,
- AI narrative section copy, rather than only an AI *setup* draft,
- the branded PDF uploaded to Cloudinary so there is a link to send, not just
  a blob in the database,
- filing the finished proposal onto the client's record, so it appears on
  Client 360 next to everything else we've sent them,
- pushing an opportunity into Smart 1 Suite when a proposal goes to a client.

Also here:

- Stores every quote (full builder state + generated PDF) in a database
  (SQLite by default; set DATABASE_URL for Render Postgres).
- Assigns quote numbers in the Q-10200 series (own counter, IO-style).
- Word (.docx) export; the latest PDF is archived with the quote.
- "What's still missing before this can become an IO" gap lists.
- Multiple target areas per campaign, via `hub.target_areas` — the same shape
  the IO builder reads, so a three-location client survives the hand-off.

The AI helper routes (audience estimate, business description, landing-page
review, ZIP-radius lookup) are served **here**. They used to be fetched from
`IO_API_BASE + "/sales/builder/api/..."`, a path that exists on neither this
app nor the IO app, so every one of those buttons had been silently falling
back to its placeholder. The IO conversion still calls the IO API for the
order number and the two IO PDFs, because those belong to the IO.
"""

import json
import os
import re
import secrets
import threading
import logging
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path

import requests
from dotenv import load_dotenv
from flask import Flask, jsonify, render_template, request, send_file
from flask_cors import CORS
from sqlalchemy import (Column, DateTime, Integer, LargeBinary, String, Text,
                        func, or_)
from sqlalchemy.orm import declarative_base
from sqlalchemy.orm import object_session as _sa_object_session

# ---- PDF (reportlab, same stack as the IO app) ----
from reportlab.lib import colors as rl_colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (Image as RLImage, KeepTogether, Paragraph,
                                SimpleDocTemplate, Spacer, Table, TableStyle)
from xml.sax.saxutils import escape as xml_escape

# ---- Word export ----
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from hub.extensions import create_all_metadata, session_factory, shared_engine
from hub.webargs import clamp_int

# Shared services. Per the migration rule in CLAUDE.md: this module was being
# edited anyway, so it moves onto the shared implementations rather than
# keeping local copies of geography, prompt and Suite logic.
from hub import business_description as hub_desc
from hub import creative_needs as hub_creative
from hub import current_marketing as hub_discovery
from hub import industries as hub_industries
from hub import kpi_framework as hub_kpi
from hub import proposal_spec as hub_spec
from hub import quote_validity as hub_validity
from hub import rate_card as hub_rate_card
from hub import target_areas as hub_areas
from hub import target_map as hub_map
from hub import view_tracking as hub_views
from hub.config import settings as hub_config

# Activity logging. Guarded so this module still runs standalone, but
# inside the Hub every action is attributable — this module created client quotes,
# and an unattributable change to a client's account is one nobody can
# explain later.
try:
    from hub import audit as _hub_audit
    _audit = _hub_audit.for_module("sales_builder")
except Exception:  # noqa: BLE001
    def _audit(*a, **k):  # no-op outside the Hub
        return None


load_dotenv()

BASE_DIR = Path(__file__).resolve().parent
app = Flask(__name__, template_folder=str(BASE_DIR / "templates"))
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

CORS(app, resources={r"/api/*": {"origins": "*"}}, methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"])

# =====================================================================
# Database (SQLite locally, Postgres on Render via DATABASE_URL)
# =====================================================================
# One engine for the whole Hub, from hub/extensions -- see the note in
# modules/scans/app.py. Quotes are client work: a second pool with its own
# SQLite fallback inside the container image meant "back up the database" had a
# different answer here than everywhere else.
engine = shared_engine()
SessionLocal = session_factory()
Base = declarative_base()

QUOTE_BASE = 10199  # first quote is Q-10200 (matches the IO number family)
VALID_STATUSES = ["Draft", "Sent", "Approved", "Lost", "Expired", "Converted"]


class Counter(Base):
    __tablename__ = "counters"
    name = Column(String(50), primary_key=True)
    value = Column(Integer, nullable=False, default=QUOTE_BASE)


class Quote(Base):
    __tablename__ = "quotes"
    id = Column(Integer, primary_key=True, autoincrement=True)
    quote_number = Column(String(20), unique=True, nullable=False, index=True)
    status = Column(String(20), nullable=False, default="Draft", index=True)
    client = Column(String(200), default="")
    website = Column(String(300), default="")
    industry = Column(String(100), default="")
    salesperson = Column(String(120), default="")
    data = Column(Text, default="{}")           # full builder state (JSON)
    monthly_budget = Column(Integer, default=0)
    months = Column(Integer, default=1)
    total_budget = Column(Integer, default=0)
    package = Column(String(40), default="")
    products_summary = Column(String(500), default="")
    goals_summary = Column(String(300), default="")
    geo_summary = Column(String(300), default="")
    revision = Column(Integer, default=1)
    io_number = Column(String(20), default="")
    io_client_pdf_url = Column(String(600), default="")
    io_internal_pdf_url = Column(String(600), default="")
    converted_at = Column(DateTime, nullable=True)
    pdf_blob = Column(LargeBinary, nullable=True)
    pdf_filename = Column(String(300), default="")
    pdf_generated_at = Column(DateTime, nullable=True)
    # Delivery: where the client-facing copy lives, and what it created in
    # Smart 1 Suite. The blob above is the archive; a blob cannot be emailed,
    # which is why the standalone builder uploaded to Cloudinary and this one
    # now does too.
    pdf_url = Column(String(600), default="")
    client_filed_as = Column(String(64), default="")     # hub.proposals record id
    suite_contact_id = Column(String(64), default="")
    suite_opportunity_id = Column(String(64), default="")
    delivered_at = Column(DateTime, nullable=True)
    # Who opened it, from the Hub session rather than from the state. The
    # `salesperson` column above is the *sales contact typed onto the
    # proposal* -- a field a rep fills in for the client's benefit, blank on
    # most drafts and sometimes somebody else's name entirely. The list was
    # showing that, so "who wrote this?" had no answer on the one screen the
    # question is asked.
    created_by = Column(String(160), default="")
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    updated_at = Column(DateTime, default=lambda: datetime.now(timezone.utc),
                        onupdate=lambda: datetime.now(timezone.utc))


# ---------------------------------------------------------------------------
# The client's copy: one link, what it was opened by, and whether they said yes
#
# Three tables rather than columns on `quotes`, because `create_all()` creates
# missing tables and never adds a column to an existing one -- six columns
# added to the live Postgres would be silently absent there while every local
# test passed, which is the trap `hub_user_profiles` and `cb_render_approvals`
# were each made their own table to avoid.
#
# They are also genuinely three things. A share is an address that outlives
# the revisions sent to it. A view is an event. An acceptance is a statement
# about **one specific revision**, so it is a row rather than a flag -- an
# edit does not delete what a client agreed to last week, it supersedes it,
# and both facts have to survive for the panel to say so.
# ---------------------------------------------------------------------------
class QuoteShare(Base):
    __tablename__ = "quote_shares"
    id = Column(Integer, primary_key=True, autoincrement=True)
    quote_id = Column(Integer, index=True, nullable=False)
    # One token per quote, minted once and kept. A second token would mean a
    # link already in a client's inbox stops working the day somebody presses
    # the button again.
    token = Column(String(64), unique=True, index=True, nullable=False)
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    created_by = Column(String(120), default="")
    revoked_at = Column(DateTime, nullable=True)
    # The revision the rep last deliberately sent. The page always shows the
    # current one; this is how the panel can say "you have edited it since".
    sent_revision = Column(Integer, default=1)
    sent_at = Column(DateTime, nullable=True)


class QuoteView(Base):
    __tablename__ = "quote_views"
    id = Column(Integer, primary_key=True, autoincrement=True)
    quote_id = Column(Integer, index=True, nullable=False)
    token = Column(String(64), index=True, nullable=False)
    at = Column(DateTime, default=lambda: datetime.now(timezone.utc), index=True)
    # Which revision they were reading. A count with no revision on it cannot
    # answer "did they ever open the one I sent on Tuesday".
    revision = Column(Integer, default=1)
    # A keyed digest, never an address -- see hub/view_tracking.py. Used to
    # recognise a reload inside the window and for nothing else.
    visitor = Column(String(64), default="", index=True)
    device = Column(String(20), default="")


class QuoteAcceptance(Base):
    __tablename__ = "quote_acceptances"
    id = Column(Integer, primary_key=True, autoincrement=True)
    quote_id = Column(Integer, index=True, nullable=False)
    token = Column(String(64), index=True, default="")
    revision = Column(Integer, default=1)
    at = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    name = Column(String(200), default="")
    email = Column(String(200), default="")
    visitor = Column(String(64), default="")


class Activity(Base):
    __tablename__ = "activity"
    id = Column(Integer, primary_key=True, autoincrement=True)
    quote_id = Column(Integer, index=True, nullable=True)
    icon = Column(String(10), default="•")
    text = Column(String(500), default="")
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))


# Advisory-locked, and reported rather than raised: a database slow to wake
# must not take the module offline for the life of the worker.
DB_BOOT_ERROR = create_all_metadata(Base.metadata)
if DB_BOOT_ERROR:
    logger.error("sales_builder: table creation reported: %s", DB_BOOT_ERROR)


# Columns added to `quotes` after the table was first created, as
# (column, type). create_all() creates missing TABLES, never missing columns,
# so the live database -- which has quotes in it already -- would keep working
# right up until the first query mentioning a delivery column.
_LATE_QUOTE_COLUMNS = [
    ("pdf_url", "VARCHAR(600)"),
    ("client_filed_as", "VARCHAR(64)"),
    ("suite_contact_id", "VARCHAR(64)"),
    ("suite_opportunity_id", "VARCHAR(64)"),
    ("delivered_at", "TIMESTAMP"),
    ("created_by", "VARCHAR(160)"),
]


def _add_missing_columns() -> None:
    """Add the columns above, asking first which ones are actually missing.

    This used to fire all five unconditionally and swallow the failures, on the
    reasoning that "already exists" is the normal case after the first boot. It
    is worse than that: all five are declared on the Quote model, so
    create_all() puts them on a fresh database too and ALL FIVE fail on EVERY
    boot, on every database, including the first. Two gunicorn workers, so ten
    Postgres ERROR lines per deploy that mean nothing -- and a log that always
    carries ten fake errors is a log nobody finds the real one in. CI surfaced
    it: the Postgres service prints its log at the end of every run.

    modules/image_picker/models.py had the identical bug and the identical fix.
    ADD COLUMN IF NOT EXISTS would be shorter, and is what modules/sites_admin
    uses -- but that module talks to Postgres directly, and this one shares the
    Hub engine, which is SQLite in local development, where the answer is
    `near "EXISTS": syntax error`. The inspector is the same answer on both.
    """
    from sqlalchemy import inspect as _inspect, text as _text
    try:
        have = {c["name"] for c in _inspect(engine).get_columns("quotes")}
    except Exception:                                   # noqa: BLE001
        return                                          # no table: nothing to alter
    for column, coltype in _LATE_QUOTE_COLUMNS:
        if column in have:
            continue
        try:
            with engine.begin() as conn:
                conn.execute(_text(
                    f"ALTER TABLE quotes ADD COLUMN {column} {coltype}"))
        except Exception:                               # noqa: BLE001
            pass                                        # raced by the other worker


if not DB_BOOT_ERROR:
    _add_missing_columns()

_COUNTER_LOCK = threading.Lock()


def next_quote_number(db):
    """Allocate the next quote number: Q-10200, Q-10201, ... (own counter)."""
    with _COUNTER_LOCK:
        row = db.get(Counter, "quote_number")
        if row is None:
            row = Counter(name="quote_number", value=QUOTE_BASE)
            db.add(row)
        row.value = int(row.value) + 1
        db.flush()
        return f"Q-{row.value}"


def log_activity(db, quote_id, icon, text):
    db.add(Activity(quote_id=quote_id, icon=icon, text=text[:490]))


# =====================================================================
# Gap check — what a finished IO still needs that a proposal may not have
# (mirrors the required pieces of the IO flow; rule-based so it works
#  with or without AI)
# =====================================================================
def compute_gaps(state):
    gaps = []
    s = state or {}

    def blank(v):
        return v in (None, "", [], {})

    if not hub_areas.normalize(s.get("targetAreas")) and not hub_areas.from_legacy(s):
        gaps.append({"key": "geo", "label": "At least one target area to run in"})
    if blank(s.get("clientContactName")) and blank(s.get("clientContactEmail")):
        gaps.append({"key": "contact", "label": "Client contact name, email, and phone"})
    if blank(s.get("startDate")):
        gaps.append({"key": "dates", "label": "Exact campaign start date (and end date or ongoing)"})
    # Video and audio each need their own answer before the plan is priced --
    # a Connected TV buy with no spot is a launch date nobody can hit. Display
    # is still the one blanket question below it.
    gaps.extend(hub_creative.gaps(s))
    gaps.extend(hub_discovery.gaps(s))
    if blank(s.get("creativeSource")):
        gaps.append({"key": "creative", "label": "Creative assets — client provides, or Smart 1 builds (creative fee)"})
    tp = s.get("trackingPlan") or {}
    if blank(tp.get("primaryConversion")) or blank(tp.get("ga4")):
        gaps.append({"key": "tracking", "label": "Tracking plan (primary conversion, GA4, call tracking, verifier)"})
    if blank(s.get("landingUrl")):
        gaps.append({"key": "landing", "label": "Final landing page URL"})
    if blank(s.get("salesContact")) or blank(s.get("salesEmail")):
        gaps.append({"key": "sales", "label": "Smart 1 sales contact name and email"})
    if blank(s.get("managementFeeAck")):
        gaps.append({"key": "fees", "label": "Management fee confirmation (and creative fee if Smart 1 builds)"})
    return gaps


# Guardrails (subset of the IO builder's rules, checked server-side too)
def compute_guardrails(state):
    """What is wrong with this plan, in the words the wizard already uses.

    The minimum is the **card's own, per product**, through
    `rate_card.minimum_for()`. This held every SEARCH ENGINE MARKETING line to
    a flat $1,500, and the card sells paid search from $400 — so it warned
    about perfectly valid $500 and $1,000 search buys, quoting a figure that
    is not that product's minimum, while saying nothing at all about a $500
    Connected TV line whose real floor is $1,500 and which the IO refuses.
    Wrong in both directions, each screen internally consistent.

    `guardrailsJs()` in the wizard was fixed for exactly this and its comment
    says "paid search is now $400 in one place and every document reads it
    there" — which was true of the screen a rep edits on and false of this,
    the reading that rides on the quote payload into the proposals list, the
    dashboard nudges and `ioDataPayload()`'s `guardrailWarnings`. So the
    insertion order carried the stale answer while the wizard showed the
    right one. The two halves produce identical strings now, and
    test_campaign_cost.py runs them against each other.
    """
    warns = []
    s = state or {}
    items = s.get("items") or []
    budget = float(s.get("budget") or 0)
    months = int(s.get("months") or 1)
    for it in items:
        dollars = float(it.get("dollars") or 0)
        # A one-time build is not a monthly buy and has no monthly floor to be
        # under -- judging it against one blocks a website-only proposal.
        if (it.get("basis") or "monthly") == "one_time":
            continue
        minimum = hub_rate_card.minimum_for(it.get("product") or "",
                                            it.get("category") or "")
        if 0 < dollars < minimum:
            name = it.get("product") or it.get("category") or "That product"
            warns.append(f"{name} at {_money(dollars)}/mo is below its "
                         f"{_money(minimum)} monthly minimum — the IO will "
                         f"not accept it.")
    if budget > 0 and len(items) > 4 and budget / max(len(items), 1) < 750:
        warns.append("Budget is split across many products — consider fewer products for impact.")
    if months < 3:
        warns.append("Campaign term under 3 months — most Smart 1 programs need 3+ months to optimize.")
    if (s.get("creativeSource") or "").lower().startswith("smart 1") and not s.get("creativeFee"):
        warns.append("Smart 1 is building creative but no creative fee is set.")
    return warns


# =====================================================================
# Serialization
# =====================================================================
def _client_key(name, url):
    """Which client record this quote belongs to, joined the shared way.

    hub/client_key.py is the single place that decides what a client is --
    d:<domain> where there is a URL, n:<name-slug> where there is not. Matching
    on the name alone is what attributed "Acme" to whichever of Acme Plumbing,
    Acme Roofing and Acme Electric came out of a dict first.
    """
    try:
        from hub import client_key as hub_client_key
        return hub_client_key.client_key(name or "", url or "")
    except Exception:                                       # noqa: BLE001
        return ""


def _signed_in_as() -> str:
    """Who is using the builder, from the Hub session.

    `AuthGuard` puts the display name on the WSGI environ before Flask sees
    the request, which is the only place it exists for a dispatcher-mounted
    module -- this app has no Hub app context and no `flask.g` of its own.
    Empty is a real answer: the module runs standalone in the tests and
    behind the shared password in an emergency, and "" reads as "not
    recorded" on the list rather than as somebody's name.
    """
    try:
        return str(request.environ.get("s1hub.user") or "").strip()[:160]
    except Exception:                                   # noqa: BLE001
        return ""


_UNSET = object()


def _sent_at_of(q):
    """When this quote was last deliberately sent, or None.

    Read through the session that loaded the quote -- one small indexed query
    -- so every caller of quote_json() gets the right answer without nine call
    sites having to remember. `list_quotes` prefetches instead, because one
    query per row on a 300-row list is the N+1 this would otherwise be.
    """
    try:
        db = _sa_object_session(q)
        if db is None:
            return None
        share = (db.query(QuoteShare)
                 .filter(QuoteShare.quote_id == q.id)
                 .order_by(QuoteShare.id.desc()).first())
        return share.sent_at if share else None
    except Exception:                                   # noqa: BLE001
        return None


def _sent_at_map(db, quotes) -> dict:
    """{quote_id: sent_at} for a page of quotes, in one query."""
    ids = [q.id for q in quotes]
    if not ids:
        return {}
    out = {}
    try:
        for share in (db.query(QuoteShare)
                      .filter(QuoteShare.quote_id.in_(ids))
                      .order_by(QuoteShare.id.asc()).all()):
            out[share.quote_id] = share.sent_at
    except Exception:                                   # noqa: BLE001
        return {}
    return out


def _state_of(q) -> dict:
    """A quote's saved answers, or an empty dict. Never raises: half this
    module reads the blob to answer one question about it, and a malformed one
    must cost that answer rather than the page."""
    try:
        state = json.loads(q.data or "{}")
    except (TypeError, ValueError):
        return {}
    return state if isinstance(state, dict) else {}


def _quote_window(q, state, sent_at=_UNSET):
    """How long this quote's pricing stands. hub/quote_validity.py has the rules."""
    if sent_at is _UNSET:
        sent_at = _sent_at_of(q)
    try:
        return hub_validity.window(q.status or "", sent_at=sent_at,
                                   created_at=q.created_at, state=state)
    except Exception:                                   # noqa: BLE001
        # A window that cannot be computed must not cost the quote: every
        # screen reads `applies`, and False is "no window here", which is what
        # every quote had before this existed.
        return {"applies": False, "measured": False, "expired": False}


def _validity_block(win):
    """What a screen is handed. The sentence travels with the dates, so no
    caller can render a date without the words that explain it."""
    out = dict(win or {})
    out["note"] = hub_validity.staff_note(win)
    out["client_note"] = hub_validity.client_note(win)
    return out


def quote_json(q, include_data=False, sent_at=_UNSET):
    state = {}
    try:
        state = json.loads(q.data or "{}")
    except Exception:
        state = {}
    win = _quote_window(q, state, sent_at)
    out = {
        "id": q.id,
        "quote_number": q.quote_number,
        "status": q.status,
        # What a screen shows. `status` stays exactly as stored, because that
        # is what a status change writes back and what the picker sets; this
        # is the same fact with the clock applied. Two fields rather than one
        # so nothing can round-trip a derived value into the column.
        "shown_status": (hub_validity.EXPIRED_STATUS if win.get("expired")
                         else q.status),
        "client": q.client,
        "website": q.website,
        "industry": q.industry,
        "salesperson": q.salesperson,
        # Two different questions, and the list was answering the wrong one.
        # `salesperson` is the sales contact typed onto the proposal for the
        # client's benefit; `created_by` is who opened it here.
        "created_by": getattr(q, "created_by", "") or "",
        "monthly_budget": q.monthly_budget,
        "months": q.months,
        "total_budget": q.total_budget,
        "package": q.package,
        "products_summary": q.products_summary,
        "goals_summary": q.goals_summary,
        "geo_summary": q.geo_summary,
        # Derived on read, never stored. CLAUDE.md: a stored key goes stale the
        # moment a client is renamed in Knack, and create_all() would not add
        # the column to the live Postgres anyway -- every local test would pass
        # against a column that was silently absent in production.
        "client_key": _client_key(q.client, q.website),
        "revision": q.revision,
        # How long the pricing stands, and whether it still does. Derived on
        # read and never stored: there are two gunicorn workers, so a status
        # written by whichever one ran a sweep is one the other disagrees
        # with -- and a stored "Expired" would survive an extension, leaving a
        # quote reading as dead on the one screen a rep would go to revive it.
        # hub/quote_validity.py owns the rules.
        "validity": _validity_block(win),
        # How far into the wizard this quote was left. Read from the state
        # blob rather than a column of its own: create_all() adds no column
        # to an existing table, so one here would be silently absent on the
        # live Postgres with every local test green. It is what lets the list
        # say a draft is half-finished instead of only that it is a draft.
        "step": (state.get("_step") if isinstance(state.get("_step"), int) else 0),
        "io_number": q.io_number,
        "io_client_pdf_url": q.io_client_pdf_url,
        "io_internal_pdf_url": q.io_internal_pdf_url,
        "has_pdf": bool(q.pdf_blob),
        "pdf_filename": q.pdf_filename,
        "created_at": q.created_at.isoformat() if q.created_at else "",
        "updated_at": q.updated_at.isoformat() if q.updated_at else "",
        "converted_at": q.converted_at.isoformat() if q.converted_at else "",
        "gaps": compute_gaps(state),
        "growth": growth_options(state),
        "guardrails": compute_guardrails(state),
        # Computed here rather than stored, exactly like `growth` above: it is
        # derived from the KPIs and the media mix, and a stale copy of it is a
        # measurement framework that no longer matches the plan under it.
        "kpi_framework": hub_kpi.framework(state),
        "target_areas": campaign_areas(state),
        "targets_of_interest": targets_of_interest(state),
        "zip_exceptions": hub_areas.zip_exceptions(campaign_areas(state)),
        "creative": hub_creative.evaluate(state),
        "suggestions": hub_discovery.suggestions(state),
        "pdf_url": q.pdf_url or "",
        "suite_opportunity_id": q.suite_opportunity_id or "",
        "delivered_at": q.delivered_at.isoformat() if q.delivered_at else "",
    }
    if include_data:
        out["data"] = state
        # The plan as the PDF and the Word export draw it, so the builder's
        # preview shows the delivery figures rather than carrying a fourth
        # copy of the arithmetic. Only on a single quote: the list does not
        # draw this table, and computing it per row there is work nothing
        # reads.
        out["media_plan"] = media_plan_rows(state)
        # What the campaign costs, and the investment summary as the PDF and
        # the Word export draw it. Shipped for the same reason the media plan
        # is: the preview had its own copy of this arithmetic, which summed
        # every line including a one-time production as though it recurred.
        out["campaign_cost"] = campaign_cost(state)
        out["investment"] = investment_lines(state, q)
    return out


def targets_of_interest(state) -> list[dict]:
    """The competitors, venues and places this campaign is going after.

    The audience step offered "Competitor conquesting" as a tick box and
    nothing anywhere asked which competitors -- so the proposal promised to
    target a client's rivals without naming one and the insertion order
    arrived with the same two words on it. Named rows only: an empty name is
    a row somebody started and abandoned, and it must not reach a client
    document as a blank line.

    A row with no address is kept. Conquesting by brand and browsing
    behaviour needs no location at all, and dropping those would mean the
    list is only ever as long as somebody's patience for looking up street
    addresses. `fenceable` says which of the two each row can do, so nothing
    downstream has to guess -- and nothing here invents an address from a
    name, the rule modules/ads_builder/logo.py works to.
    """
    rows = []
    for raw in (state or {}).get("targetsOfInterest") or []:
        if not isinstance(raw, dict):
            continue
        name = str(raw.get("name") or "").strip()
        if not name:
            continue
        address = str(raw.get("address") or "").strip()
        rows.append({
            "kind": str(raw.get("kind") or "competitor").strip() or "competitor",
            "name": name[:200],
            "address": address[:200],
            "note": str(raw.get("note") or "").strip()[:200],
            "fenceable": bool(address),
        })
    return rows


def campaign_areas(state):
    """Every target area on a quote, old records included.

    Quotes saved before multi-area targeting have their geography in
    `geoType` / `geo` / `radius` only, so `from_legacy` reads those rather
    than a migration rewriting rows nobody has re-opened. A quote that has
    both keeps the explicit list: it was edited more recently than it was
    imported.
    """
    state = state or {}
    return hub_areas.normalize(state.get("targetAreas")) or hub_areas.from_legacy(state)


def summarize_into(q, state):
    """Pull list/summary columns out of the saved state."""
    q.client = str(state.get("client") or "")[:200]
    q.website = str(state.get("url") or "")[:300]
    q.industry = str(state.get("industry") or "")[:100]
    q.salesperson = str(state.get("salesContact") or "")[:120]
    q.months = int(state.get("months") or 1)
    # The plan is the number. This used to read the selected package or the
    # budget the rep typed on the Budget step -- what the client *asked for* --
    # so a plan edited down to $5,750 still put $8,000 on the cover, in the
    # investment summary and on the dashboard's pipeline, while the media mix
    # table and the insertion order both said $5,750. See campaign_cost().
    cost = campaign_cost(state)
    sel = state.get("selectedPackage") or {}
    if cost["has_plan"]:
        q.monthly_budget = int(round(cost["recurring"]))
        q.total_budget = int(round(cost["campaign"]))
    else:
        # No plan yet: the ask is the only figure there is, and a quote at the
        # Budget step still has to read sensibly on the list.
        q.monthly_budget = int(sel.get("monthly") or state.get("budget") or 0)
        q.total_budget = int(sel.get("total") or (q.monthly_budget * q.months))
    q.package = str(sel.get("name") or "")[:40]
    items = state.get("items") or []
    q.products_summary = " · ".join([str(i.get("product") or "") for i in items])[:500]
    q.goals_summary = ", ".join(state.get("objectives") or [])[:300]
    # The list column says how many places this campaign runs in, not just the
    # first one -- "Carmel showroom + 3 more" is the difference between a
    # single-location quote and a four-rooftop one at a glance.
    areas = campaign_areas(state)
    q.geo_summary = (hub_areas.summary(areas) or str(state.get("geo") or ""))[:300]


# =====================================================================
# Routes — core
# =====================================================================
def _io_api_base():
    """Where the IO builder's own routes live.

    It is mounted inside the Hub at /tools/io, so the default is that mount:
    same origin, same login, no cold start. IO_API_BASE still overrides it if
    the IO ever moves back to its own service.

    This used to default to the external Render URL, and the callers then
    appended "/sales/builder/api/..." to it -- a path that exists on neither
    app. Every conversion call 404'd.
    """
    return os.getenv("IO_API_BASE", "/tools/io").rstrip("/")


@app.get("/health")
def health():
    return jsonify({
        "status": "ok",
        # engine.url, not a module-level _db_url -- that name was never
        # defined, so this route raised NameError and answered 500 to every
        # "is the Sales Builder up?" check ever made against it.
        "database": engine.url.get_backend_name(),
        "suite": _suite_status(),
        "openai_configured": bool(os.getenv("OPENAI_API_KEY")),
        "io_api_base": _io_api_base(),
    })


@app.get("/")
def index():
    return render_template("index.html")


@app.get("/api/config")
def api_config():
    return jsonify({
        # Through _io_api_base(), not a second os.getenv with its own default.
        # This read still carried the OLD external default -- the very one the
        # docstring on that function describes as the bug -- and it is the read
        # the browser actually uses: index.html seeds CFG with "/tools/io" and
        # then assigns whatever this route returns over the top. So /health
        # reported the mount while every proposal-to-IO conversion posted to a
        # different Render service: a cold start, a different login, and a
        # "The IO API did not return an order number." in the conversion log
        # with nothing saying where it had gone. One reader, the rule this
        # codebase applies to rate cards, client keys and source labels alike.
        "io_api_base": _io_api_base(),
        # Now mounted inside the Hub rather than iframed from Render, so it
        # shares the login and can reach the client registry. The external URL
        # still works as an override if the standalone app is ever needed.
        "io_app_url": os.getenv("IO_APP_URL", "/tools/io/?embed=1"),
        "ai_enabled": bool(os.getenv("OPENAI_API_KEY")),
        # Served rather than mirrored. Target areas and the creative
        # classifier each carry a JavaScript copy of a server rule, and each
        # needs a test proving the two halves still agree; that cost is paid
        # twice already, so the sizes and the markup rule come down the wire
        # and the browser renders what it is given -- the choice Social
        # Planner made about its calendar, for the same reason.
        "creative_sizes": _creative_sizes(),
        "rate_rules": hub_rate_card.rate_rules_for_js(),
        # The cover's industry-trends block, served rather than mirrored: the
        # PDF and the Word export read hub_spec.industry_trends() directly,
        # and the preview reads this -- one table, three renderers.
        "industry_trends": {"note": hub_spec.TRENDS_NOTE,
                            "general": hub_spec.GENERAL_TRENDS,
                            "industries": hub_spec.INDUSTRY_TRENDS},
        # The KPI choices the Measurement step offers — the IO builder's own
        # benchmark table, served for the same reason the sizes above are:
        # the per-product half already reaches the page as kpi_framework.rows
        # on the quote payload, and this is the static half, so a fresh quote
        # with no goals and no plan still gets a full choice list rather than
        # an empty pill row over an "Add a KPI…" box.
        "kpi_choices": hub_kpi.choices(),
    })


def _creative_sizes() -> dict:
    """The spec-kit sizes each gated medium needs, for the creative step.

    The exact answer is per product and lives in
    `hub.creative_needs.required_units()`, which is what travels to the
    insertion order. This is the same kit read one level up, so the wizard can
    say "here is what we will be asking the client for" before a product has
    been picked. A medium the kit maps nothing for is left out rather than
    sent as an empty list a screen would render as "no creative needed".
    """
    try:
        from hub import creative_specs
    except Exception:                                   # noqa: BLE001
        return {}
    # One representative product per medium, only to reach the kit's own
    # channel mapping. Naming products here rather than channels keeps this
    # honest: if the mapping changes, this follows it.
    probes = {
        "video": ("Connected TV - Targeted", "OTT"),
        "audio": ("Programmatic - Targeted", "DIGITAL RADIO"),
        "display": ("Select Tactics - Comes with Retargeting", "DATA TARGETED DISPLAY"),
        "retargeting": ("Website Retargeting", "RETARGETING"),
    }
    out = {}
    for medium, (product, category) in probes.items():
        # The same sentence the proposal prints, from the same function.
        # Built here from a flat list of pixel sizes, the audio row read
        # "300x250" -- the optional companion banner presented as the whole
        # requirement, on the screen that asks whether the client has the
        # creative. Somebody sends a banner and no spot.
        probe = {"months": 1, "items": [{"product": product, "category": category,
                                         "dollars": 1}]}
        line = hub_creative.units_line(probe, medium)
        detail = hub_creative.required_units(probe, medium)
        if detail["measured"]:
            out[medium] = {"line": line,
                           "sizes": [sz for unit in detail["units"]
                                     for sz in unit["sizes"]],
                           "source": detail.get("source") or ""}
    return out


# ---- Quotes CRUD ----
@app.post("/api/quotes")
def create_quote():
    body = request.get_json(force=True) or {}
    state = body.get("data") or {}
    db = SessionLocal()
    try:
        q = Quote(quote_number=next_quote_number(db),
                  created_by=_signed_in_as(),
                  status=body.get("status") or "Draft",
                  data=json.dumps(state, ensure_ascii=False))
        summarize_into(q, state)
        db.add(q)
        db.flush()
        log_activity(db, q.id, "🆕", f"{q.quote_number} created — {q.client or 'new quote'}")
        db.commit()
        return jsonify({"ok": True, "quote": quote_json(q, include_data=True)})
    finally:
        db.close()


@app.get("/api/quotes")
def list_quotes():
    qstr = (request.args.get("q") or "").strip().lower()
    status = (request.args.get("status") or "").strip()
    limit = clamp_int(request.args.get("limit"), 50, 1, 500)
    db = SessionLocal()
    try:
        query = db.query(Quote)
        if status and status != "all":
            query = query.filter(Quote.status == status)
        if qstr:
            like = f"%{qstr}%"
            query = query.filter(or_(
                func.lower(Quote.quote_number).like(like),
                func.lower(Quote.client).like(like),
                func.lower(Quote.website).like(like),
                func.lower(Quote.products_summary).like(like),
                func.lower(Quote.io_number).like(like),
            ))
        rows = query.order_by(Quote.updated_at.desc()).limit(limit).all()
        sent = _sent_at_map(db, rows)
        return jsonify({"ok": True,
                        "quotes": [quote_json(r, sent_at=sent.get(r.id))
                                   for r in rows]})
    finally:
        db.close()


@app.get("/api/uploaded-proposals")
def list_uploaded_proposals():
    """Proposals written outside this tool and uploaded to a client record.

    Kept as a separate call rather than merged into /api/quotes because the
    two are genuinely different objects: a quote has a number, a revision and
    a status this tool owns, while an uploaded proposal is a file with a date.
    Flattening them into one row shape would mean inventing values for half
    the columns, and an invented "Draft" on a document nobody here drafted is
    exactly the sort of confident wrong answer to avoid. The front end merges
    them for display and keeps them labelled.
    """
    limit = clamp_int(request.args.get("limit"), 100, 1, 500)
    try:
        from hub import proposals as hub_proposals
    except Exception:                                   # noqa: BLE001
        # Standalone, outside the Hub. An empty list is the honest answer.
        return jsonify({"ok": True, "proposals": [], "available": False})
    try:
        rows = hub_proposals.all_proposals(limit=limit,
                                           q=(request.args.get("q") or "").strip())
    except Exception as exc:                            # noqa: BLE001
        logger.warning("uploaded proposals unavailable: %s", exc)
        return jsonify({"ok": True, "proposals": [], "available": False,
                        "note": "Uploaded proposals could not be read."})
    return jsonify({"ok": True, "proposals": rows, "available": True})


@app.get("/api/quotes/<int:qid>")
def get_quote(qid):
    db = SessionLocal()
    try:
        q = db.get(Quote, qid)
        if not q:
            return jsonify({"ok": False, "error": "Quote not found"}), 404
        return jsonify({"ok": True, "quote": quote_json(q, include_data=True)})
    finally:
        db.close()


@app.put("/api/quotes/<int:qid>")
def update_quote(qid):
    body = request.get_json(force=True) or {}
    db = SessionLocal()
    try:
        q = db.get(Quote, qid)
        if not q:
            return jsonify({"ok": False, "error": "Quote not found"}), 404
        if "data" in body:
            state = body.get("data") or {}
            q.data = json.dumps(state, ensure_ascii=False)
            summarize_into(q, state)
            q.revision = (q.revision or 1) + (1 if body.get("bump_revision") else 0)
        if "status" in body:
            new_status = body["status"]
            if new_status not in VALID_STATUSES:
                return jsonify({"ok": False, "error": "Invalid status"}), 400
            if new_status != q.status:
                log_activity(db, q.id, {"Sent": "📤", "Approved": "✅", "Lost": "❌",
                                        "Converted": "🔁", "Expired": "⏰"}.get(new_status, "•"),
                             f"{q.quote_number} → {new_status} — {q.client}")
            q.status = new_status
        db.commit()
        return jsonify({"ok": True, "quote": quote_json(q, include_data=True)})
    finally:
        db.close()


@app.delete("/api/quotes/<int:qid>")
def delete_quote(qid):
    """Delete a draft quote.

    Drafts only, and the check is here rather than only in the UI. Every other
    status means the document left the building: a Sent quote is one a client
    has read, an Approved one is what an insertion order was agreed from, and
    a Converted one is the paper trail behind a live campaign. Losing any of
    those to a mis-click is unrecoverable — there is no undo and no backup of
    a single row — so the only thing this removes is a draft nobody outside
    the office has seen.

    A draft that was delivered before being set back to draft still has its
    filed copy on the client's record; that is a different object in
    hub.proposals and is deliberately left alone. The response says so rather
    than implying everything went.
    """
    db = SessionLocal()
    try:
        q = db.get(Quote, qid)
        if not q:
            return jsonify({"ok": False, "error": "Quote not found"}), 404
        if q.status != "Draft":
            return jsonify({"ok": False,
                            "error": f"{q.quote_number} is {q.status.lower()}, not a draft. "
                                     f"Only drafts can be deleted — mark it Lost or Expired "
                                     f"instead so the history is kept."}), 409

        number, client = q.quote_number, q.client
        filed = bool(q.client_filed_as)
        # No FK and no cascade on Activity.quote_id, so its rows would be left
        # pointing at an id that no longer exists and the feed would render
        # entries for a quote nobody can open.
        db.query(Activity).filter(Activity.quote_id == q.id).delete(synchronize_session=False)
        db.delete(q)
        # Logged with no quote_id, because the quote it refers to is gone. The
        # feed should still be able to say that it was deleted.
        log_activity(db, None, "🗑️", f"{number} deleted — {client or 'no client'}")
        db.commit()
        _audit("quote_deleted", client=client, quote=number)
        return jsonify({"ok": True, "deleted": number, "filed_copy_kept": filed})
    finally:
        db.close()


@app.post("/api/quotes/<int:qid>/duplicate")
def duplicate_quote(qid):
    db = SessionLocal()
    try:
        src = db.get(Quote, qid)
        if not src:
            return jsonify({"ok": False, "error": "Quote not found"}), 404
        state = json.loads(src.data or "{}")
        # New quotes start clean of decision/IO fields
        for k in ("startDate", "ioPayload",):
            state.pop(k, None)
        # A duplicate is a new proposal, so it is credited to whoever made
        # it rather than to whoever wrote the one it was copied from.
        q = Quote(quote_number=next_quote_number(db), status="Draft",
                  created_by=_signed_in_as(),
                  data=json.dumps(state, ensure_ascii=False))
        summarize_into(q, state)
        db.add(q)
        db.flush()
        log_activity(db, q.id, "📋", f"{q.quote_number} duplicated from {src.quote_number} — {q.client}")
        db.commit()
        return jsonify({"ok": True, "quote": quote_json(q, include_data=True)})
    finally:
        db.close()


@app.post("/api/quotes/<int:qid>/converted")
def mark_converted(qid):
    """Called by the convert wizard after the IO API has issued the order
    number and generated the IO PDFs. Links everything back to the quote."""
    body = request.get_json(force=True) or {}
    db = SessionLocal()
    try:
        q = db.get(Quote, qid)
        if not q:
            return jsonify({"ok": False, "error": "Quote not found"}), 404
        q.io_number = str(body.get("io_number") or "")[:20]
        q.io_client_pdf_url = str(body.get("client_pdf_url") or "")[:600]
        q.io_internal_pdf_url = str(body.get("internal_pdf_url") or "")[:600]
        q.status = "Converted"
        q.converted_at = datetime.now(timezone.utc)
        if "data" in body:
            q.data = json.dumps(body["data"], ensure_ascii=False)
            summarize_into(q, body["data"])
        log_activity(db, q.id, "🔁",
                     f"{q.quote_number} → IO #{q.io_number} — {q.client}"
                     + (", sent to Smart 1 Suite" if body.get("submitted") else ""))
        db.commit()
        return jsonify({"ok": True, "quote": quote_json(q)})
    finally:
        db.close()


def _pipeline_signals() -> dict:
    """What needs chasing, read by the Hub dashboard and by this one."""
    try:
        from hub import sales_status
        return sales_status.scoreboard()
    except Exception as exc:                            # noqa: BLE001
        return {"measured": False,
                "error": f"The pipeline could not be read ({type(exc).__name__})."}


# ---- Dashboard ----
@app.get("/api/dashboard")
def dashboard():
    db = SessionLocal()
    try:
        rows = db.query(Quote).all()
        now = datetime.now(timezone.utc)

        def aware(dt):
            if dt is None:
                return now
            return dt if dt.tzinfo else dt.replace(tzinfo=timezone.utc)

        open_rows = [r for r in rows if r.status in ("Draft", "Sent")]
        sent_rows = [r for r in rows if r.status == "Sent"]
        approved = [r for r in rows if r.status == "Approved"]
        conv_month = [r for r in rows if r.status == "Converted" and r.converted_at
                      and aware(r.converted_at).month == now.month and aware(r.converted_at).year == now.year]
        decided_90 = [r for r in rows if r.status in ("Approved", "Converted", "Lost")
                      and (now - aware(r.updated_at)).days <= 90]
        won_90 = [r for r in decided_90 if r.status in ("Approved", "Converted")]
        avg_days_out = round(sum((now - aware(r.updated_at)).days for r in sent_rows) / len(sent_rows)) if sent_rows else 0

        acts = db.query(Activity).order_by(Activity.created_at.desc()).limit(12).all()
        stale = [r for r in sent_rows if (now - aware(r.updated_at)).days >= 7]
        nudges = []
        sent_map = _sent_at_map(db, stale[:3])
        for r in stale[:3]:
            # Whether the pricing still stands changes what the follow-up is.
            # "Chase this" and "this one needs re-quoting before they can say
            # yes" are different jobs, and the second one has a client sitting
            # in front of a page that will not let them accept.
            win = _quote_window(r, _state_of(r), sent_map.get(r.id))
            nudges.append({"tag": "FOLLOW-UP NUDGE", "text": f"{r.client} ({r.quote_number}) was sent "
                          f"{(now - aware(r.updated_at)).days} days ago with no response. "
                          + ("The pricing has expired — re-send it to quote at "
                             "current rates." if win.get("expired")
                             else "Consider a follow-up.")})
        for r in approved[:3]:
            g = compute_gaps(json.loads(r.data or "{}"))
            nudges.append({"tag": "READY TO CONVERT", "text": f"{r.client} ({r.quote_number}) is approved. "
                          + (f"{len(g)} answers are still needed before the IO is complete." if g
                             else "Everything needed for the IO is already on file.")})
        for r in [x for x in rows if x.status == "Draft"][:5]:
            for w in compute_guardrails(json.loads(r.data or "{}"))[:1]:
                nudges.append({"tag": "PRICING GUARDRAIL", "text": f"{r.client or r.quote_number}: {w}"})

        return jsonify({"ok": True, "stats": {
            "open_count": len(open_rows),
            "open_monthly": sum(r.monthly_budget or 0 for r in open_rows),
            "awaiting_count": len(sent_rows),
            "awaiting_avg_days": avg_days_out,
            "approved_count": len(approved),
            "converted_month_count": len(conv_month),
            "converted_month_monthly": sum(r.monthly_budget or 0 for r in conv_month),
            "win_rate_90d": round(100 * len(won_90) / len(decided_90)) if decided_90 else 0,
        },
            # The same five signals the Hub dashboard's pipeline card draws,
            # from the same reading. Two screens answering "what needs
            # chasing" separately is how they come to disagree in front of
            # the same rep -- the /api/db/structure versus /api/integrity
            # trap. Never allowed to cost this dashboard: a failure here is
            # `measured: False` and the panel says so.
            "pipeline": _pipeline_signals(),
            "activity": [{"icon": a.icon, "text": a.text,
                         "when": a.created_at.isoformat() if a.created_at else ""} for a in acts],
            "nudges": nudges[:6]})
    finally:
        db.close()


# =====================================================================
# Branded proposal PDF
# =====================================================================
NAVY = rl_colors.HexColor("#14284b")
BLUE = rl_colors.HexColor("#1f63ae")
GOLD = rl_colors.HexColor("#e5a323")
LINE = rl_colors.HexColor("#d5dee9")
SOFT = rl_colors.HexColor("#eef3f8")
MUTED = rl_colors.HexColor("#53657a")
GOLD = rl_colors.HexColor("#e5a323")


def _sec_header(title: str, number: int, st_h2) -> list:
    """A numbered section header: gold number chip, navy title, ruled under.

    Every section used to open with a bare bold line, which on a
    fourteen-section document made the Executive Summary and the ZIP appendix
    read with identical weight. The number is computed over the ENABLED
    sections at build time rather than stored, so hiding a section renumbers
    the ones after it — a stored number would say "07" on the fifth heading
    the day somebody hid two sections, on a document a client reads. The
    preview draws the same treatment from the same rule
    (test_proposal_spec.py holds the two together).
    """
    chip = ParagraphStyle("secno", parent=st_h2, alignment=TA_CENTER,
                          textColor=rl_colors.white, spaceBefore=0, spaceAfter=0)
    ttl = ParagraphStyle("sect", parent=st_h2, spaceBefore=0, spaceAfter=0)
    t = Table([[Paragraph(f"<b>{number:02d}</b>", chip),
                Paragraph(f"<b>{xml_escape(title)}</b>", ttl)]],
              colWidths=[0.42 * inch, 6.98 * inch])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, 0), GOLD),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LINEBELOW", (0, 0), (-1, 0), 0.9, LINE),
        ("LEFTPADDING", (1, 0), (1, 0), 9),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4)]))
    space_before = getattr(st_h2, "spaceBefore", 12) or 12
    space_after = getattr(st_h2, "spaceAfter", 5) or 5
    return [Spacer(1, space_before), t, Spacer(1, space_after)]


def _trends_flowables(state, st_body, st_small) -> list:
    """The cover's standing block: what is moving in this client's category,
    how Smart 1 answers it, and how budgets like this are usually crafted.

    Read from `hub_spec.industry_trends()` — a table, never a model call,
    because this is the first thing a client reads (the reasoning is on the
    table itself). Always present: an unknown industry gets the general
    entry rather than a hole on the first page.
    """
    t = hub_spec.industry_trends(state.get("industry"))
    lead = ParagraphStyle("trlead", parent=st_body, textColor=NAVY,
                          spaceBefore=6, spaceAfter=2)
    out = [Paragraph(f"<b>What is happening in "
                     f"{xml_escape(t['industry'] if t['matched'] else 'digital marketing')}</b>",
                     lead)]
    for line in t["trends"]:
        out += _body_flowables("• " + line, st_body)
    out.append(Paragraph("<b>How Smart 1 helps</b>", lead))
    out.append(Paragraph(xml_escape(t["help"]), st_body))
    out.append(Paragraph("<b>How budgets like this are usually crafted</b>", lead))
    out.append(Paragraph(xml_escape(t["budget"]), st_body))
    out.append(Paragraph(xml_escape(t["note"]), st_small))
    return out


# The one tag generated copy is allowed to carry. `clean_ai_text` normalises
# every other form of emphasis away and leaves this, so escaping the string
# and then putting these two back is the whole of "bold is fine": reportlab
# reads <b> natively, and nothing else in the copy can reach it as markup.
_BOLD_OPEN, _BOLD_CLOSE = xml_escape("<b>"), xml_escape("</b>")


def _p(text, style):
    body = xml_escape(str(text or "")).replace("\n", "<br/>")
    body = body.replace(_BOLD_OPEN, "<b>").replace(_BOLD_CLOSE, "</b>")
    return Paragraph(body, style)


def _inline(text) -> str:
    """One line of cleaned copy, escaped, with the one allowed tag put back."""
    body = xml_escape(str(text or ""))
    return body.replace(_BOLD_OPEN, "<b>").replace(_BOLD_CLOSE, "</b>")


def show_map(state) -> bool:
    """Whether this proposal carries a coverage map.

    Default on, and switched off per proposal the same way a generated table
    is -- one flag on the areas section, read by the preview, the PDF and the
    Word export, rather than three screens each deciding for themselves.
    """
    for sec in (state or {}).get("sections") or []:
        if sec.get("kind") == "areas":
            return sec.get("showMap") is not False and sec.get("enabled", True)
    # No outline on this quote yet -- which is every quote at the moment the
    # target areas are being described, three steps before the document
    # exists. Answering False there meant the map panel on the areas screen
    # said "left out of this proposal" about a proposal with no sections in
    # it at all. Only an explicit no is a no.
    return True


def campaign_map(state):
    """The campaign's map, and what it does and does not show. Never raises."""
    if not show_map(state):
        return None, {"reason": "The map is left out of this proposal.",
                      "plotted": [], "not_plotted": [], "measured": False}
    try:
        return hub_map.render(campaign_areas(state))
    except Exception as exc:                            # noqa: BLE001
        logger.warning("target map failed: %s", exc)
        return None, {"reason": f"The map could not be drawn ({exc}).",
                      "plotted": [], "not_plotted": [], "measured": False}


def _docx_body(d, text) -> None:
    """Section copy in Word, with a bulleted list as a bulleted list.

    Word has a List Bullet style and python-docx will raise if the template
    does not carry it, so the fallback writes the bullet character itself --
    a list that looks slightly hand-made beats an export that 500s, and beats
    the run-on sentence this replaced.
    """
    for block in hub_spec.blocks(text):
        rows = ([block] if block["kind"] == "para" else block["items"])
        for row in rows:
            para = None
            if block["kind"] == "list":
                try:
                    para = d.add_paragraph(style="List Bullet")
                except Exception:                   # noqa: BLE001
                    para = d.add_paragraph()
                    para.add_run("\u2022  ")
            if para is None:
                para = d.add_paragraph()
            for piece, bold in row["runs"]:
                run = para.add_run(piece)
                run.font.bold = bold


def _docx_map(d, png: bytes, meta: dict) -> None:
    """The same map the PDF carries, at the same place in the document."""
    try:
        # Bounded on both axes for the reason `_map_flowables` gives: width
        # alone lets a tall crop fill a page on its own.
        from PIL import Image as PILImage
        with PILImage.open(BytesIO(png)) as probe:
            pw, ph = probe.size
        scale = min(6.4 / max(1, pw), 4.0 / max(1, ph))
        d.add_picture(BytesIO(png), width=Inches(pw * scale),
                      height=Inches(ph * scale))
    except Exception as exc:                        # noqa: BLE001
        logger.warning("map could not be placed in the Word export: %s", exc)
        return
    covered = [row["label"] for row in meta.get("not_plotted") or []
               if row.get("kind") in hub_map.NOT_DRAWN]
    if covered:
        note = d.add_paragraph()
        run = note.add_run("Also covered, and not drawn on the map: "
                           + ", ".join(covered) + ".")
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x53, 0x65, 0x7A)


# How big the map is allowed to be on the page. Half the text column's height
# is plenty to recognise where a campaign runs, and it leaves the section's
# copy and its table on the same page as the picture far more often.
MAP_MAX_W = 7.4 * inch
MAP_MAX_H = 4.3 * inch


# The map's share of the row when the ZIP column sits beside it: two thirds,
# with the targeted ZIPs in the remaining third. Bounded per area and in
# total, because a reportlab table row does not split across pages -- an
# unbounded DMA list beside the picture is the one flowable the frame cannot
# place, and it takes the whole PDF with it. The full list still prints in
# ZIP Codes Targeted, which is what the truncation note points at.
MAP_ZIP_W = 4.85 * inch
ZIP_COL_W = 2.45 * inch
ZIP_COL_PER_AREA = 24
ZIP_COL_MAX_AREAS = 4


def _zip_column_rows(state) -> list:
    """(label, zips) per complete area, for the column beside the map."""
    rows = []
    for area in campaign_areas(state):
        zips = hub_areas.zip_list(area.get("zips"))
        if zips:
            rows.append((hub_areas.label(area), zips))
    return rows


def _zip_column_flowables(zip_rows, small_style) -> list:
    """The 1/3 column: which ZIPs the campaign targets, per area, bounded."""
    lead = ParagraphStyle("zc", parent=small_style, textColor=NAVY)
    out = [Paragraph("<b>ZIP Codes targeted</b>", lead)]
    for label, zips in zip_rows[:ZIP_COL_MAX_AREAS]:
        shown = ", ".join(zips[:ZIP_COL_PER_AREA])
        more = len(zips) - ZIP_COL_PER_AREA
        out.append(Paragraph(
            f"<b>{xml_escape(label)}</b> · {len(zips)} ZIP "
            f"Code{'' if len(zips) == 1 else 's'}<br/>{xml_escape(shown)}"
            + (f" <i>+{more} more</i>" if more > 0 else ""), small_style))
    left_out = len(zip_rows) - ZIP_COL_MAX_AREAS
    if left_out > 0:
        out.append(Paragraph(f"<i>+{left_out} more area"
                             f"{'' if left_out == 1 else 's'}</i>", small_style))
    if left_out > 0 or any(len(z) > ZIP_COL_PER_AREA for _, z in zip_rows):
        out.append(Paragraph("The complete list is in ZIP Codes Targeted, at "
                             "the back of this proposal.", small_style))
    return out


def _map_flowables(png: bytes, meta: dict, small_style, zip_rows=None) -> list:
    """The map at two thirds of the row, the targeted ZIPs in the third
    beside it, and what the picture leaves out named underneath.

    The ZIP column exists because the picture and the list answer the same
    question at two zoom levels -- where the campaign runs, and exactly which
    ZIPs that means -- and a client reads them together. A campaign with no
    ZIPs anywhere (a DMA, a state, a national buy) gets the full-width map it
    always had: a two-thirds picture beside an empty column is a layout
    holding space for nothing.

    The caveat line is deliberately short and factual. It exists because a
    map showing two rings on a campaign that also covers a whole DMA is a
    picture of *part* of the buy, and a client reading it as the whole buy is
    the confidently wrong answer this codebase keeps having to undo. What is
    never printed here is the internal half -- a geocoder that could not find
    a spelling is the rep's to fix, and belongs on the builder's screen.
    """
    from PIL import Image as PILImage           # already a dependency
    try:
        with PILImage.open(BytesIO(png)) as im:
            width, height = im.size
    except Exception:                           # noqa: BLE001
        return []
    # Bounded on BOTH axes, not just scaled to the column width.
    #
    # A three-rooftop campaign running north-south cropped to 511x606, and at
    # the column width that is a picture 8.8 inches tall: most of page two, a
    # half-empty page one above it, and one slightly taller campaign away from
    # a flowable the page frame cannot place at all. `target_map` keeps the
    # crop landscape now and this is the guarantee at the other end, because a
    # map that cannot be placed takes the whole PDF with it.
    zip_rows = zip_rows or []
    max_w = MAP_ZIP_W if zip_rows else MAP_MAX_W
    scale = min(max_w / max(1, width), MAP_MAX_H / max(1, height))
    image = RLImage(BytesIO(png), width=width * scale, height=height * scale)
    image.hAlign = "CENTER"
    if zip_rows:
        row = Table([[image, _zip_column_flowables(zip_rows, small_style)]],
                    colWidths=[MAP_ZIP_W + 0.1 * inch, ZIP_COL_W])
        row.setStyle(TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("BACKGROUND", (1, 0), (1, 0), SOFT),
            ("LEFTPADDING", (1, 0), (1, 0), 8),
            ("RIGHTPADDING", (1, 0), (1, 0), 8),
            ("TOPPADDING", (1, 0), (1, 0), 6),
            ("BOTTOMPADDING", (1, 0), (1, 0), 6),
            ("LEFTPADDING", (0, 0), (0, 0), 0)]))
        block = [row]
    else:
        block = [image]
    covered = [row["label"] for row in meta.get("not_plotted") or []
               if row.get("kind") in hub_map.NOT_DRAWN]
    if covered:
        block.append(_p("Also covered, and not drawn on the map: "
                        + ", ".join(covered) + ".", small_style))
    # The caption is about the picture directly above it. Orphaned onto the
    # next page it is a sentence about nothing, on a document a client reads.
    return [Spacer(1, 4), KeepTogether(block), Spacer(1, 4)]


def _body_flowables(text, style) -> list:
    """Section copy as paragraphs and real bullet lists.

    A bulleted list used to arrive here as one string with the bullet
    characters inside it, so reportlab set it as a paragraph and the client
    read "we will target three areas: • Carmel • Fishers • Noblesville" as a
    sentence. `proposal_spec.blocks()` is the one place that decides what is a
    list -- the preview and the Word export read the same function, so the
    three renderers cannot disagree about it.
    """
    out = []
    for block in hub_spec.blocks(text):
        if block["kind"] == "para":
            out.append(Paragraph(_inline(block["text"]).replace("\n", "<br/>"), style))
            continue
        bullet = ParagraphStyle(
            f"bul{id(block)}", parent=style, leftIndent=style.fontSize * 1.6,
            bulletIndent=style.fontSize * 0.5, spaceAfter=max(1, style.spaceAfter - 2))
        for item in block["items"]:
            out.append(Paragraph(_inline(item["text"]), bullet, bulletText="\u2022"))
        if out:
            out.append(Spacer(1, 3))
    return out


# Sections whose generated table a rep may leave out or replace by hand.
# `text` and `cover` carry no table, so offering them the control would make
# the button mean nothing.
TABLE_KINDS = ("areas", "reach", "friction", "channels", "creative",
               "kpis", "mediaplan", "packages", "roi", "growth", "zips")


def section_table(sec) -> dict:
    """What to render under one section: the generated table, an edit, or none.

    Three answers, and the document has to honour all three. The builder can
    leave a table out (the copy above it still prints) or replace it with one
    edited by hand -- for the case the generator cannot cover, a row naming a
    location under NDA or a KPI the client asked us to drop. Read here rather
    than in each of the twelve `kind` branches so a section added later
    cannot quietly ignore the setting, which is how the creative check came
    to be missing from the Preview panel's label map.

        {"show": bool, "rows": [[...]] or None, "head": [...] or None}
    """
    sec = sec or {}
    if sec.get("kind") not in TABLE_KINDS:
        return {"show": True, "rows": None, "head": None}
    if sec.get("showTable") is False:
        return {"show": False, "rows": None, "head": None}
    edit = sec.get("tableEdit")
    if isinstance(edit, dict) and isinstance(edit.get("rows"), list) and edit["rows"]:
        head = [str(h) for h in (edit.get("head") or [])]
        rows = [[str(c) for c in row] for row in edit["rows"] if isinstance(row, list)]
        return {"show": True, "rows": rows, "head": head}
    return {"show": True, "rows": None, "head": None}


def _sell_rate(item):
    """What one media line is quoted at, or None where there is no rate.

    The card's rate is the buy-side number. Every CPM and CPV line is quoted
    at a multiple of it -- 2x to start, and a rep's own number once they set
    one -- and everything downstream of the plan has to read the same figure,
    or the delivery table promises impressions the campaign cannot buy. A
    management-fee, flat-fee or custom-quote line has nothing to multiply and
    reports None rather than a plausible number.
    """
    item = item or {}
    rate_type = item.get("rate") or item.get("rate_type")
    if not hub_rate_card.is_marked_up(rate_type):
        return None
    try:
        own = float(item.get("sellRate") or 0)
    except (TypeError, ValueError):
        own = 0.0
    if own > 0:
        return round(own, 2)
    return hub_rate_card.sell_rate(item.get("rateValue") or item.get("rate_value"),
                                   rate_type)


def _money(n):
    try:
        return "$" + f"{float(n):,.0f}"
    except Exception:
        return str(n)


def growth_options(state) -> dict:
    """Two grounded routes to a bigger plan, for the foot of the proposal.

    A client who says yes usually asks "and what would more look like?", and
    the answer has always been improvised on a call. It belongs on the
    document -- but only if the numbers are defensible, so neither route
    invents one:

      * **Raise a budget.** The difference between what each line is quoted at
        and what the Accelerated package already prices it at. That figure is
        on the proposal two sections above; this just names the gap.

      * **Add a product.** Whatever discovery said they are missing (the
        We Suggest They Should list), priced at the rate card's own monthly
        minimum for that product -- the smallest honest number, not a guess at
        what they might spend.

    Every row is editable on the proposal. `custom` rows are ones the rep
    added or re-priced, and they are returned untouched.
    """
    state = state or {}
    months = max(1, int(state.get("months") or 1))
    items = state.get("items") or []
    on_plan = {str(i.get("product") or "").strip().lower() for i in items}

    # --- raising what is already there ---
    # The Accelerated package is built at 150% of budget and is already on the
    # document, so the uplift per line is arithmetic rather than opinion.
    increases = []
    accelerated = None
    for pkg in state.get("packages") or []:
        if str(pkg.get("name") or "").lower().startswith("acceler"):
            accelerated = pkg
            break
    by_product = {}
    for line in ((accelerated or {}).get("lines") or []):
        by_product[str(line.get("product") or "").strip().lower()] = line
    for item in items:
        now = float(item.get("dollars") or 0)
        target = float((by_product.get(str(item.get("product") or "").strip().lower())
                        or {}).get("dollars") or 0)
        if target > now > 0:
            increases.append({
                "product": item.get("product", ""),
                "category": item.get("category", ""),
                "current": round(now, 2),
                "suggested": round(target, 2),
                "uplift": round(target - now, 2),
                "campaign": round((target - now) * months, 2),
            })

    # --- adding what is not there ---
    additions = []
    seen = set()
    try:
        from hub import current_marketing as hub_marketing
        wanted = hub_marketing.suggestions(state)
    except Exception:                                       # noqa: BLE001
        wanted = []
    for suggestion in wanted:
        for name in suggestion.get("products") or []:
            key = str(name).strip().lower()
            if not key or key in on_plan or key in seen:
                continue
            seen.add(key)
            card = hub_rate_card.find(name) or {}
            # No card entry means no defensible price, so it is offered as
            # something to quote rather than with a number nobody can stand
            # behind.
            minimum = (hub_rate_card.minimum_for(name, card.get("category", ""))
                       if card else 0)
            additions.append({
                "product": card.get("product") or name,
                "category": card.get("category", ""),
                "why": suggestion.get("title", ""),
                "suggested": round(float(minimum), 2) if minimum else 0,
                "campaign": round(float(minimum) * months, 2) if minimum else 0,
                "rate": card.get("listed_rate", ""),
                "quoted": bool(minimum),
            })

    edits = state.get("growthEdits") or {}
    for row in increases:
        override = edits.get("inc:" + row["product"])
        if override is not None:
            try:
                row["suggested"] = round(float(override), 2)
                row["uplift"] = round(row["suggested"] - row["current"], 2)
                row["campaign"] = round(row["uplift"] * months, 2)
                row["custom"] = True
            except (TypeError, ValueError):
                pass
    for row in additions:
        override = edits.get("add:" + row["product"])
        if override is not None:
            try:
                row["suggested"] = round(float(override), 2)
                row["campaign"] = round(row["suggested"] * months, 2)
                row["quoted"] = row["suggested"] > 0
                row["custom"] = True
            except (TypeError, ValueError):
                pass

    dropped = set(state.get("growthDropped") or [])
    increases = [r for r in increases if ("inc:" + r["product"]) not in dropped]
    additions = [r for r in additions if ("add:" + r["product"]) not in dropped]

    return {
        "increases": increases,
        "additions": additions,
        "months": months,
        "increase_monthly": round(sum(r["uplift"] for r in increases), 2),
        "addition_monthly": round(sum(r["suggested"] for r in additions), 2),
        "any": bool(increases or additions),
    }


def investment_lines(state, q):
    """Platform, media and one-time production, kept apart.

    The specification requires recurring SaaS fees to be separated from media
    spend and one-time setup. Blending them is how a client comes to believe
    the platform stops costing money when they pause a campaign.
    """
    state = state or {}
    months = max(1, int(getattr(q, "months", 0) or state.get("months") or 1))
    # From the plan, not from the quote's summary column and not from a naked
    # sum of every line: that sum counted a one-time production line as if it
    # were charged every month, so a $1,500 shoot added $1,500 to the monthly
    # figure and $9,000 to a six-month campaign.
    cost = campaign_cost(state)
    monthly_media = cost["recurring"] if cost["has_plan"] else \
        float(getattr(q, "monthly_budget", 0) or 0)

    lines = []
    suite = state.get("suiteTier") or {}
    if suite.get("include") is not False:
        # The tier's own name and specs come from the specification; only the
        # price is the rep's to move. Reading `specs` off whatever the wizard
        # saved left the line as "Smart 1 Suite — Smarter ()" the moment a
        # tier was picked, because the wizard stores the choice rather than a
        # copy of the tier.
        listed = hub_spec.suggested_tier(monthly_media)
        if suite.get("name"):
            listed = next((t for t in hub_spec.SAAS_TIERS
                           if t["name"] == suite["name"]), listed)
        try:
            quoted = float(suite["monthly"]) if suite.get("monthly") is not None \
                else float(listed.get("monthly") or 0)
        except (TypeError, ValueError):
            quoted = float(listed.get("monthly") or 0)
        lines.append({"label": f"Smart 1 Suite — {listed['name']} ({listed.get('specs', '')})",
                      "amount": quoted,
                      "recurs": "Monthly", "kind": "saas",
                      # Printed nowhere on the client document. It is here so
                      # the internal copy and the IO can tell a negotiated
                      # price from the list one -- a discount nobody recorded
                      # is a discount nobody can renew.
                      "listed": float(listed.get("monthly") or 0),
                      "adjusted": abs(quoted - float(listed.get("monthly") or 0)) > 0.001})
    # Consulting & Strategy rides beside the licence, never inside it: it is
    # Smart 1's time working the Suite products with the client every month,
    # priced from the hours the wizard estimated and editable by the rep. Like
    # the Suite it is recurring platform work rather than media, so pausing
    # the campaign does not stop it -- which is exactly why it is its own
    # line. The hours ride into the label so the client can see what the
    # figure is made of.
    consulting = state.get("consulting") or {}
    if consulting.get("include"):
        try:
            c_listed = float(consulting.get("listed") or 0)
        except (TypeError, ValueError):
            c_listed = 0.0
        try:
            c_amount = float(consulting["monthly"]) \
                if consulting.get("monthly") is not None else c_listed
        except (TypeError, ValueError):
            c_amount = c_listed
        try:
            c_hours = float(consulting.get("hours") or 0)
        except (TypeError, ValueError):
            c_hours = 0.0
        label = "Consulting & Strategy — monthly Suite coaching and campaign strategy"
        if c_hours > 0:
            label += f" (~{c_hours:g} hrs/mo)"
        lines.append({"label": label, "amount": c_amount,
                      "recurs": "Monthly", "kind": "consulting",
                      "listed": c_listed,
                      "adjusted": abs(c_amount - c_listed) > 0.001})
    lines.append({"label": "Campaign media & services",
                  "amount": round(monthly_media, 2),
                  "recurs": "Monthly", "kind": "media"})
    # One-time lines on the plan are their own rows rather than being folded
    # into a monthly figure. Named individually, because "$3,250 one-time" on
    # a document a client signs is a number they are entitled to see the
    # parts of.
    for row in cost["one_time_lines"]:
        lines.append({"label": row["label"], "amount": row["amount"],
                      "recurs": "One-time", "kind": "setup"})

    creative = hub_creative.evaluate(state)
    for row in creative["media"]:
        if row["answer"] == hub_creative.CLIENT_PAYS and row["fee"]:
            lines.append({"label": f"{row['label'].split(' (')[0]} creative production",
                          "amount": float(row["fee"]), "recurs": "One-time",
                          "kind": "setup"})
    try:
        extra = float(str(state.get("creativeFee") or "0").replace("$", "").replace(",", ""))
    except ValueError:
        extra = 0.0
    if extra:
        lines.append({"label": "Display creative production", "amount": extra,
                      "recurs": "One-time", "kind": "setup"})

    recurring = sum(l["amount"] for l in lines if l["recurs"] == "Monthly")
    one_time = sum(l["amount"] for l in lines if l["recurs"] == "One-time")
    # The campaign total is the plan's own arithmetic plus what sits beside
    # it, never `recurring * months`: a line bought for three months of a
    # six-month flight is not charged for six, and the plan's one-time rows
    # are already counted inside cost["campaign"], so adding them again from
    # `lines` would bill a video shoot twice.
    plan_one_time = sum(r["amount"] for r in cost["one_time_lines"])
    campaign_media = (cost["campaign"] if cost["has_plan"]
                      else monthly_media * months)
    # The licence and the consulting retainer both recur for the whole term,
    # so both multiply by months the way the licence always has -- one bucket,
    # or the campaign total quietly stops covering the newer of the two lines.
    platform_monthly = sum(l["amount"] for l in lines
                           if l["recurs"] == "Monthly"
                           and l["kind"] in ("saas", "consulting"))
    campaign_total = (campaign_media + platform_monthly * months
                      + (one_time - plan_one_time))
    return {"lines": lines, "recurring_monthly": round(recurring, 2),
            "one_time": round(one_time, 2),
            "first_month": round(recurring + one_time, 2),
            "campaign_total": round(campaign_total, 2),
            # What the total-campaign row says it includes. Computed here and
            # read by the PDF, the preview and the review screen alike --
            # "including licensing" over a total that also carries consulting
            # is a label understating the number directly beside it.
            "includes": ("licensing & consulting"
                         if any(l["kind"] == "consulting" for l in lines)
                         else "licensing"),
            "months": months}


def _head_style_rows():
    """The navy-header table style every generated section uses."""
    return [("BACKGROUND", (0, 0), (-1, 0), NAVY),
            ("TEXTCOLOR", (0, 0), (-1, 0), rl_colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.4, LINE),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5)]


def _head_style():
    return TableStyle(_head_style_rows())


# What each medium does in the customer journey, and what it is judged on.
# Awareness channels measured on clicks is the single most common way a good
# campaign gets called a failure.
_CHANNEL_ROLE = {
    hub_creative.VIDEO: ("Top of funnel — builds awareness and trust on the "
                         "screens the household already watches",
                         "Completed views and completion rate"),
    hub_creative.AUDIO: ("Mid funnel — reaches the 79% of digital audio time "
                         "spent with no screen at all",
                         "Reach, frequency and listen-through"),
    hub_creative.SOCIAL: ("Mid funnel — consideration and re-engagement in feed",
                          "Engaged sessions and cost per lead"),
    hub_creative.DISPLAY: ("Full funnel — precision targeting and retargeting "
                           "against the people already in market",
                           "Click-through rate and cost per action"),
    hub_creative.OTHER: ("Supports the campaign", "Cost per action"),
}


def channel_lines(state) -> list:
    """The plan's lines that are actually channels.

    A production line and a management fee are not channels, and listing them
    here put "Video Production — top of funnel, builds awareness and trust on
    the screens the household already watches" and "Management Fee — supports
    the campaign" in a table headed Recommended Channel Strategy, on a
    document a client reads. A one-time line is a cost of making the creative;
    a line with no rate and no gated medium is a fee. Neither runs anywhere.
    """
    out = []
    for item in state.get("items") or []:
        if str(item.get("basis") or "monthly") == "one_time":
            continue
        if (_sell_rate(item) is None
                and hub_creative.medium_of(item) == hub_creative.OTHER):
            continue
        out.append(item)
    return out


def _channel_role(item):
    medium = hub_creative.medium_of(item)
    category = str(item.get("category") or "").upper()
    if "SEARCH ENGINE MARKETING" in category or "PAY PER CLICK" in category:
        return ("Bottom of funnel — captures demand that is already searching",
                "Cost per lead and conversion rate")
    if "SEARCH ENGINE OPTIMIZATION" in category:
        return ("Compounding — organic and AI-answer visibility",
                "Ranked terms, map-pack visibility and organic leads")
    if "RETARGETING" in category:
        return ("Bottom of funnel — brings back the people who already came",
                "Return visits and cost per conversion")
    return _CHANNEL_ROLE.get(medium, _CHANNEL_ROLE[hub_creative.OTHER])


def _creative_phrase(row):
    """What the creative gate decided, phrased for a client-facing document."""
    if row["answer"] == hub_creative.HAS:
        return "Supplied by the client"
    if row["answer"] == hub_creative.CLIENT_PAYS:
        return f"Produced by Smart 1 — {_money(row['fee'])} one-time"
    if row["answer"] == hub_creative.COMP:
        return "Produced by Smart 1 at no charge"
    return "To be confirmed before launch"


# Roughly how much a section costs in vertical space, so the scale is chosen
# from what the document actually contains rather than from a page count we
# would only learn after building it once.
_SECTION_WEIGHT = {"mediaplan": 3, "packages": 3, "roi": 4, "areas": 2,
                   "channels": 3, "creative": 2, "friction": 2,
                   "reach": 1, "kpis": 1, "cover": 0}

TYPE_SCALE_MIN = 0.82
TYPE_SCALE_FULL_UNDER = 38


def _type_scale(state) -> float:
    """How far to shrink the type so a full proposal stays compact.

    Estimated from the content: characters of prose, plus a weight per
    generated table and a row-count for the two that grow with the campaign.
    Building the PDF twice to measure the real page count would be exact, but
    reportlab gives no page count until it has built, and a rep waiting twice
    as long for a PDF would notice that far more than the half-point of type.
    """
    state = state or {}
    sections = [sec for sec in state.get("sections") or [] if sec.get("enabled", True)]
    prose = sum(len(str(sec.get("body") or "")) for sec in sections)
    weight = sum(_SECTION_WEIGHT.get(sec.get("kind"), 1) for sec in sections)
    # The two tables whose height follows the campaign rather than the outline.
    weight += len(state.get("items") or []) + len(campaign_areas(state))

    load = prose / 900.0 + weight

    # A freshly seeded proposal — the full outline, its default copy, one
    # product, one area — measures about 34. That is the baseline a normal
    # proposal sits at, so it must scale at 1.0: "lower the fonts when
    # necessary" means when there is more than usual in the document, not
    # always. Shrinking every proposal by default would just make the type
    # smaller and prove nothing.
    if load <= TYPE_SCALE_FULL_UNDER:
        return 1.0
    # Ease down to the floor rather than stepping: two proposals of similar
    # length should not come out at visibly different type sizes.
    return round(max(TYPE_SCALE_MIN, 1.0 - (load - TYPE_SCALE_FULL_UNDER) * 0.009), 3)


def build_proposal_pdf(q, state, sent_at=_UNSET):
    title = f"S1M Proposal - {q.quote_number} - {q.client or 'Client'}"
    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter, rightMargin=0.55 * inch, leftMargin=0.55 * inch,
                            topMargin=0.55 * inch, bottomMargin=0.6 * inch, title=title)
    ss = getSampleStyleSheet()
    # Type scales down as the document grows, so a proposal with a lot in it
    # stays compact rather than sprawling. Bounded at 0.82: below that the
    # body copy stops being comfortable to read on paper, and a proposal
    # nobody wants to read is worse than one that runs an extra page.
    scale = _type_scale(state)

    def sized(size, leading):
        return round(size * scale, 1), round(leading * scale, 1)

    ts, tl = sized(21, 25)
    ss_, sl = sized(10.5, 13)
    hs, hl = sized(13, 16)
    bs, bl = sized(9.5, 13)
    ms, ml = sized(8, 10.5)

    st_title = ParagraphStyle("T", parent=ss["Title"], textColor=NAVY, fontSize=ts, leading=tl, alignment=TA_CENTER, spaceAfter=2)
    st_sub = ParagraphStyle("Sub", parent=ss["BodyText"], textColor=MUTED, fontSize=ss_, alignment=TA_CENTER, spaceAfter=round(10 * scale))
    st_h2 = ParagraphStyle("H2", parent=ss["Heading2"], textColor=NAVY, fontSize=hs, leading=hl, spaceBefore=round(13 * scale), spaceAfter=round(5 * scale))
    st_body = ParagraphStyle("B", parent=ss["BodyText"], fontSize=bs, leading=bl, spaceAfter=round(5 * scale))
    st_small = ParagraphStyle("S", parent=ss["BodyText"], fontSize=ms, leading=ml, textColor=MUTED)

    story = []
    # Branded header band
    band = Table([[Paragraph('<font color="#ffffff"><b>SMART 1 MARKETING</b></font>', ParagraphStyle("bnd", fontSize=13, leading=16, alignment=TA_CENTER))]],
                 colWidths=[7.4 * inch], rowHeights=[0.42 * inch])
    band.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), NAVY), ("VALIGN", (0, 0), (-1, -1), "MIDDLE")]))
    story += [band, Spacer(1, 14),
              Paragraph(xml_escape(f"Marketing Proposal — {q.client or 'Client'}"), st_title),
              Paragraph(xml_escape(f"Quote {q.quote_number}  ·  {datetime.now().strftime('%B %d, %Y')}"
                                   + (f"  ·  Prepared by {state.get('salesContact')}" if state.get("salesContact") else "")), st_sub)]

    # Meta table
    sel = state.get("selectedPackage") or {}
    areas = campaign_areas(state)
    meta = [["Client", q.client or ""],
            ["Website", q.website or ""],
            # One per line, the way the target-area cell beside it already
            # reads. Three goals joined by commas in a two-inch cell wraps
            # into a paragraph and stops looking like a list of three things.
            ["Campaign Goals", "\n".join(state.get("objectives") or [])],
            ["Target Area" if len(areas) < 2 else f"Target Areas ({len(areas)})",
             "\n".join(hub_areas.names(areas)) or q.geo_summary or ""],
            ["Term", f"{q.months} months"],
            # The plan's own figures, labelled with their scope. These used
            # to read the selected package first, so the cover quoted the
            # budget the client asked for while the media mix five pages
            # later totalled what was actually being bought. The licence and
            # any one-time production are added in the Investment Summary,
            # which says so; the two same-scope figures now agree exactly.
            ["Monthly campaign investment", _money(q.monthly_budget)],
            [f"Total campaign investment ({q.months} months)",
             _money(q.total_budget)]]
    # An empty row is dropped rather than printed as a labelled blank -- a
    # "Campaign Goals" row with nothing beside it reads as a question the
    # document forgot to answer, on the first thing a client sees.
    meta = [row for row in meta if str(row[1] or "").strip()]
    # Both cells are Paragraphs so both WRAP. The label used to be a bare
    # string in a 1.55in column, and reportlab does not wrap a bare string:
    # "Monthly campaign investment" ran under the value column and printed
    # over the figure -- "investme$8,050", on the cover of a real proposal.
    st_meta = ParagraphStyle("ML", parent=st_body, textColor=NAVY,
                             fontName="Helvetica-Bold", spaceAfter=0)
    meta = [[_p(row[0], st_meta), _p(row[1], st_body)] for row in meta]
    t = Table(meta, colWidths=[2.15 * inch, 5.25 * inch])
    t.setStyle(TableStyle([("BACKGROUND", (0, 0), (0, -1), SOFT),
                           ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                           ("GRID", (0, 0), (-1, -1), 0.4, LINE), ("VALIGN", (0, 0), (-1, -1), "TOP"),
                           ("LEFTPADDING", (0, 0), (-1, -1), 6), ("TOPPADDING", (0, 0), (-1, -1), 5),
                           ("BOTTOMPADDING", (0, 0), (-1, -1), 5)]))
    story += [t]

    # Sections (editable in the builder; sensible defaults if absent)
    sec_no = 0
    for sec in state.get("sections") or []:
        if not sec.get("enabled", True):
            continue
        sec_no += 1
        story += _sec_header(sec.get("title") or "", sec_no, st_h2)
        if sec.get("body"):
            story += _body_flowables(sec.get("body"), st_body)
        kind = sec.get("kind")
        # The cover always carries the industry-trends block -- what is moving
        # in this category, how Smart 1 answers it, and how budgets like this
        # are usually crafted. Standing content, not a table a rep can
        # exclude: it is the reason the rest of the document exists.
        if kind == "cover":
            story += _trends_flowables(state, st_body, st_small)
        # Excluded or hand-edited, decided once. A branch per kind is how the
        # setting comes to be honoured by eleven of the twelve.
        table_plan = section_table(sec)
        if not table_plan["show"]:
            continue
        if table_plan["rows"] is not None:
            head = table_plan["head"] or []
            rows = ([[_p(c, st_small) for c in head]] if head else []) + \
                   [[_p(c, st_small) for c in row] for row in table_plan["rows"]]
            width = max((len(r) for r in rows), default=0)
            if width:
                for row in rows:
                    row.extend([_p("", st_small)] * (width - len(row)))
                col = (7.4 * inch) / width
                et = Table(rows, colWidths=[col] * width,
                           repeatRows=1 if head else 0)
                et.setStyle(_head_style() if head
                            else TableStyle(_head_style_rows()[2:]))
                story.append(et)
            continue
        if kind == "areas":
            # The map first, then the table that names what it shows. A client
            # recognises where the campaign runs far faster than they read
            # three sentences describing it -- and the areas a map cannot draw
            # (a DMA, a state, a national buy) are in the table underneath,
            # which is why the picture never stands alone.
            png, map_meta = campaign_map(state)
            if png:
                story += _map_flowables(png, map_meta, st_small,
                                        zip_rows=_zip_column_rows(state))
            if areas:
                rows = [["Target area", "Coverage", "Est. population"]]
                for area in areas:
                    population = hub_areas.estimated_population(area)
                    zips = hub_areas.zip_list(area.get("zips"))
                    coverage = area.get("notes") or (f"{len(zips)} ZIP Codes" if zips else area["type"])
                    rows.append([_p(hub_areas.label(area), st_small),
                                 _p(coverage, st_small),
                                 # "Not measured", never a zero. An area we
                                 # could not size must not read as an area
                                 # with nobody in it.
                                 f"{population:,}" if population else "Not measured"])
                at = Table(rows, colWidths=[3.5 * inch, 2.2 * inch, 1.7 * inch], repeatRows=1)
                at.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), NAVY),
                                        ("TEXTCOLOR", (0, 0), (-1, 0), rl_colors.white),
                                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                                        ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                                        ("GRID", (0, 0), (-1, -1), 0.4, LINE),
                                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                                        ("ALIGN", (2, 1), (-1, -1), "RIGHT"),
                                        ("TOPPADDING", (0, 0), (-1, -1), 5),
                                        ("BOTTOMPADDING", (0, 0), (-1, -1), 5)]))
                story.append(at)
                if len(areas) > 1:
                    story.append(_p("Populations are shown per area. Where areas share "
                                    "ZIP Codes, the campaign totals count the shared "
                                    "ZIPs once rather than adding the overlap again.",
                                    st_small))
            targets = targets_of_interest(state)
            if targets:
                rows = [["Who we are going after", "What we want from them",
                         "How we reach them"]]
                for row in targets:
                    rows.append([
                        _p(row["name"], st_small),
                        _p(row["note"] or ("Their customers"
                                           if row["kind"] == "competitor"
                                           else "The people who go there"), st_small),
                        _p(f"Geo-fenced at {row['address']}" if row["fenceable"]
                           else "Brand and behavior — no address on file to fence",
                           st_small)])
                tt = Table(rows, colWidths=[2.1 * inch, 2.6 * inch, 2.7 * inch],
                           repeatRows=1)
                tt.setStyle(_head_style())
                story += [Spacer(1, 6), tt]
        elif kind == "reach":
            est = state.get("estimates") or {}
            if est:
                rows = [["Est. Population", "Addressable Audience", "Households", "Devices"],
                        [f"{int(est.get('pop') or 0):,}", f"{int(est.get('aud') or 0):,}",
                         f"{int(est.get('hh') or 0):,}", f"{int(est.get('dev') or 0):,}"]]
                rt = Table(rows, colWidths=[1.85 * inch] * 4)
                rt.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), NAVY), ("TEXTCOLOR", (0, 0), (-1, 0), rl_colors.white),
                                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"), ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                                        ("ALIGN", (0, 0), (-1, -1), "CENTER"), ("GRID", (0, 0), (-1, -1), 0.4, LINE),
                                        ("TOPPADDING", (0, 0), (-1, -1), 5), ("BOTTOMPADDING", (0, 0), (-1, -1), 5)]))
                story.append(rt)
            auds = state.get("audiences") or []
            if auds:
                story += _body_flowables(
                    hub_spec.bullets(auds, "Audience layers"), st_body)
        elif kind == "mediaplan":
            # One reading of the plan, shared with the Word export and the
            # builder's preview -- the three had already drifted, and the
            # delivery column is the reason it matters: three copies of that
            # arithmetic is three answers to what a line buys.
            plan = media_plan_rows(state)
            if plan["rows"]:
                rows = [[*plan["columns"][:4], f"Total ({plan['months']} mo)",
                         plan["columns"][5]]]
                for r in plan["rows"]:
                    rows.append([_p(r["product"], st_small),
                                 _p(r["category"], st_small),
                                 _p(r["rate"], st_small),
                                 (r["monthly_label"] if r["monthly_label"]
                                  else _money(r["monthly"])),
                                 _money(r["campaign"]),
                                 _p(r["delivery"], st_small)])
                # The totals a client adds up themselves if they are not
                # printed, and gets a different answer from the investment
                # summary when a one-time line is in the plan.
                rows.append(["Campaign total", "", "",
                             _money(plan["monthly_total"]),
                             _money(plan["campaign_total"]), ""])
                mt = Table(rows, colWidths=[1.85 * inch, 1.35 * inch, 1.0 * inch,
                                            0.95 * inch, 1.0 * inch, 1.25 * inch],
                           repeatRows=1)
                mt.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), NAVY), ("TEXTCOLOR", (0, 0), (-1, 0), rl_colors.white),
                                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"), ("FONTSIZE", (0, 0), (-1, -1), 8),
                                        ("GRID", (0, 0), (-1, -1), 0.4, LINE), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                                        ("ALIGN", (3, 1), (4, -1), "RIGHT"), ("TOPPADDING", (0, 0), (-1, -1), 5),
                                        ("BACKGROUND", (0, len(rows) - 1), (-1, len(rows) - 1), SOFT),
                                        ("FONTNAME", (0, len(rows) - 1), (-1, len(rows) - 1), "Helvetica-Bold"),
                                        ("BOTTOMPADDING", (0, 0), (-1, -1), 5)]))
                story.append(mt)
                # An estimate printed with no words around it reads as a
                # guarantee, which is what the ROI section was rebuilt to undo.
                if plan["note"]:
                    story += [Spacer(1, 4), Paragraph(plan["note"], st_small)]
        elif kind == "packages":
            # Recurring platform licensing, recurring media and one-time
            # production, shown apart. A client reading one blended number
            # cannot tell what stops if they pause the media.
            invest = investment_lines(state, q)
            rows = [["", "Amount", "Recurs"]]
            for line in invest["lines"]:
                rows.append([_p(line["label"], st_small), _money(line["amount"]),
                             _p(line["recurs"], st_small)])
            rows.append(["Total first month", _money(invest["first_month"]), ""])
            # Named as including the licence, because the cover and the media
            # plan both print a campaign total that deliberately does not.
            rows.append([f"Total campaign ({q.months} mo, including "
                         f"{invest.get('includes') or 'licensing'})",
                         _money(invest["campaign_total"]), ""])
            it = Table(rows, colWidths=[4.2 * inch, 1.6 * inch, 1.6 * inch], repeatRows=1)
            style = _head_style_rows()
            for offset in (2, 1):
                style.append(("BACKGROUND", (0, len(rows) - offset),
                              (-1, len(rows) - offset), SOFT))
                style.append(("FONTNAME", (0, len(rows) - offset),
                              (-1, len(rows) - offset), "Helvetica-Bold"))
            it.setStyle(TableStyle(style))
            story += [it, Spacer(1, 8)]

            pkgs = state.get("packages") or []
            if pkgs:
                selname = (state.get("selectedPackage") or {}).get("name")
                head = ["", *[p.get("name", "") + (" ★" if p.get("name") == selname else "") for p in pkgs]]
                rows = [head,
                        ["Monthly", *[_money(p.get("monthly")) for p in pkgs]],
                        [f"Total ({q.months} mo)", *[_money(p.get("total")) for p in pkgs]],
                        ["Est. impressions / mo tier", *[f"{int(p.get('impr') or 0):,}" for p in pkgs]]]
                pt = Table(rows, colWidths=[1.9 * inch] + [1.83 * inch] * len(pkgs))
                style = [("BACKGROUND", (0, 0), (-1, 0), NAVY), ("TEXTCOLOR", (0, 0), (-1, 0), rl_colors.white),
                         ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"), ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                         ("GRID", (0, 0), (-1, -1), 0.4, LINE), ("ALIGN", (1, 0), (-1, -1), "CENTER"),
                         ("FONTNAME", (0, 1), (0, -1), "Helvetica-Bold"), ("TEXTCOLOR", (0, 1), (0, -1), NAVY),
                         ("TOPPADDING", (0, 0), (-1, -1), 5), ("BOTTOMPADDING", (0, 0), (-1, -1), 5)]
                for ci, p in enumerate(pkgs):
                    if p.get("name") == selname:
                        style.append(("BACKGROUND", (ci + 1, 1), (ci + 1, -1), rl_colors.HexColor("#fff5df")))
                pt.setStyle(TableStyle(style))
                story.append(pt)
                story.append(_p("★ recommended package", st_small))
            # How long these figures stand, under the figures. An expiry the
            # client cannot see is one we cannot hold them to, and this is the
            # section it qualifies. Nothing at all on a quote that has not
            # been sent -- that one has no window yet, and a date invented for
            # it would be a promise made before the document went out.
            validity = hub_validity.client_note(
                _quote_window(q, state, sent_at=sent_at))
            if validity:
                story += [Spacer(1, 6), _p(validity, st_small)]
        elif kind == "kpis":
            # A list of things the campaign will be judged on is a list. Six
            # KPIs joined by commas is a sentence a client skims, and the
            # fourth one is the one they would have argued with.
            kpis = state.get("kpis") or []
            if kpis:
                story += _body_flowables(hub_spec.bullets(kpis), st_body)
        elif kind == "friction":
            picks = hub_discovery.suggestions(state)
            if picks:
                rows = [["We suggest they should", "Why"]]
                for pick in picks:
                    rows.append([_p(pick["title"], st_small),
                                 _p(pick["detail"], st_small)])
                ft = Table(rows, colWidths=[2.4 * inch, 5.0 * inch], repeatRows=1)
                ft.setStyle(_head_style())
                story.append(ft)
        elif kind == "channels":
            # Every channel with its role in the funnel and its KPI. The
            # specification forbids listing channels without that mapping.
            rows = [["Channel", "Role in the funnel", "Primary KPI"]]
            for item in channel_lines(state):
                role, kpi = _channel_role(item)
                rows.append([_p(item.get("product") or "", st_small),
                             _p(role, st_small), _p(kpi, st_small)])
            if len(rows) > 1:
                ct = Table(rows, colWidths=[2.3 * inch, 3.1 * inch, 2.0 * inch], repeatRows=1)
                ct.setStyle(_head_style())
                story.append(ct)
        elif kind == "creative":
            plan = hub_creative.evaluate(state)
            if plan["media"]:
                # The sizes are on the client's copy on purpose: "do you have
                # banners" has no answer until somebody says which, and the
                # client is the one being asked to send them. They come from
                # the IO's own spec kit, so what the proposal asks for and
                # what the insertion order checks are the same list.
                rows = [["Medium", "Campaign spend", "Creative", "What we need"]]
                for row in plan["media"]:
                    rows.append([_p(row["label"], st_small),
                                 _money(row["spend"]),
                                 _p(_creative_phrase(row), st_small),
                                 _p(hub_creative.units_line(state, row["medium"]),
                                    st_small)])
                kt = Table(rows, colWidths=[2.0 * inch, 1.1 * inch, 2.1 * inch,
                                            2.2 * inch], repeatRows=1)
                kt.setStyle(_head_style())
                story.append(kt)
        elif kind == "roi":
            # The KPI framework, not a table of impressions.
            #
            # "Expected Results & ROI" asks what counts as this campaign
            # working, and an impression count answers a different question:
            # it says what the money bought, not what the business gets. The
            # insertion order has carried a KPI framework all along, so the
            # two documents described one campaign two ways — the client
            # agreed to impressions and the campaign was run against KPIs.
            # `hub/kpi_framework.py` is the one description now, and the IO's
            # own copy is asserted against it.
            plan = hub_kpi.framework(state)
            if plan["measured"]:
                story += _body_flowables(
                    hub_spec.bullets([plan["primary"]], "<b>Primary KPI</b>"), st_body)
                if plan["secondary"]:
                    story += _body_flowables(
                        hub_spec.bullets(plan["secondary"], "<b>Secondary KPIs</b>"),
                        st_body)
            else:
                story.append(_p(plan["note"], st_small))
            if plan["additional_metrics"]:
                story += _body_flowables(
                    hub_spec.bullets(plan["additional_metrics"],
                                     "<b>Also tracked and reported monthly</b>"),
                    st_body)
            if plan["rows"]:
                rows = [["Product", "Key KPI", "Expected benchmark"]]
                for row in plan["rows"]:
                    rows.append([_p(row["product"], st_small),
                                 _p(row["kpi"], st_small),
                                 _p(row["expected"], st_small)])
                kt = Table(rows, colWidths=[2.9 * inch, 2.2 * inch, 2.3 * inch],
                           repeatRows=1)
                kt.setStyle(_head_style())
                story += [Spacer(1, 4), kt]
                # A range is what this inventory normally delivers. Said once,
                # in the client's own document, because a benchmark printed
                # without it reads as a number we have promised to hit.
                story.append(_p("Benchmark ranges are what this inventory "
                                "normally delivers for campaigns of this "
                                "shape. They are expectations to measure "
                                "against, not guarantees.", st_small))
            traditional = hub_discovery.roi_note(state)
            if traditional:
                story.append(_p(traditional, st_body))
        elif kind == "growth":
            growth = growth_options(state)
            if not growth["any"]:
                story.append(_p("Every product the discovery answers pointed at "
                                "is already on the plan above.", st_small))
            else:
                if growth["increases"]:
                    rows = [["Product", "Quoted", "Recommended", "More / month",
                             f"Over {growth['months']} mo"]]
                    for row in growth["increases"]:
                        rows.append([_p(row["product"], st_small),
                                     _money(row["current"]), _money(row["suggested"]),
                                     _money(row["uplift"]), _money(row["campaign"])])
                    rows.append(["Additional monthly", "", "",
                                 _money(growth["increase_monthly"]),
                                 _money(growth["increase_monthly"] * growth["months"])])
                    gt = Table(rows, colWidths=[2.6 * inch, 1.1 * inch, 1.25 * inch,
                                                1.2 * inch, 1.25 * inch], repeatRows=1)
                    style = _head_style_rows()
                    style.append(("BACKGROUND", (0, len(rows) - 1), (-1, len(rows) - 1), SOFT))
                    style.append(("FONTNAME", (0, len(rows) - 1), (-1, len(rows) - 1),
                                  "Helvetica-Bold"))
                    gt.setStyle(TableStyle(style))
                    story.append(_p("Raise a budget already on the plan", st_small))
                    story.append(gt)
                if growth["additions"]:
                    rows = [["Product", "Why", "From",
                             f"Over {growth['months']} mo"]]
                    for row in growth["additions"]:
                        rows.append([
                            _p(row["product"], st_small), _p(row["why"], st_small),
                            _money(row["suggested"]) if row["quoted"] else _p("Quoted on request", st_small),
                            _money(row["campaign"]) if row["quoted"] else _p("—", st_small)])
                    at2 = Table(rows, colWidths=[2.3 * inch, 2.8 * inch, 1.1 * inch,
                                                 1.2 * inch], repeatRows=1)
                    at2.setStyle(TableStyle(_head_style_rows()))
                    story.append(_p("Add what the discovery answers pointed at", st_small))
                    story.append(at2)
                story.append(_p("Raising a line uses the Accelerated option above; "
                                "adding one starts at our listed minimum for "
                                "that product.", st_small))
        elif kind == "zips":
            # The trafficking reference, at the back. Monospaced and small on
            # purpose: it is a list to be checked against, not read.
            # `zip_list(area)` was being handed the whole area dict, so it
            # regexed five-digit runs out of every field on it. Correct only
            # by accident, and it read the found list rather than the
            # running one -- so an area restricted to one state printed the
            # ZIPs it was restricted away from.
            rows = [["Target Area", "ZIP Codes"]]
            for area in campaign_areas(state):
                result = hub_areas.area_zips(area)
                if result["kept"]:
                    rows.append([_p(hub_areas.label(area), st_small),
                                 _p(", ".join(result["kept"]), st_small)])
            if len(rows) > 1:
                zt = Table(rows, colWidths=[2.1 * inch, 5.3 * inch], repeatRows=1)
                zt.setStyle(TableStyle(_head_style_rows()))
                story.append(zt)
            else:
                story.append(_p("No ZIP Codes were captured for this campaign. "
                                "The insertion order needs them before trafficking.",
                                st_small))
            # An exception is part of what was agreed, so it is printed where
            # the list is -- including one the system could not read, which
            # is the row that most needs somebody's eye on it.
            for rule in hub_areas.zip_exceptions(state.get("targetAreas")
                                                 or campaign_areas(state)):
                if rule["applied"]:
                    story.append(_p(f"{rule['area']} — exception: “{rule['text']}”. "
                                    f"{rule['dropped']} ZIP Code(s) excluded, "
                                    f"{rule['kept']} running.", st_small))
                else:
                    story.append(_p(f"{rule['area']} — the exception “{rule['text']}” "
                                    f"could not be read and has NOT been applied. "
                                    f"Every ZIP above is running.", st_small))


    # Footer
    story += [Spacer(1, 16),
              _p("Smart 1 Marketing  ·  smart1marketing.com  ·  This proposal is valid for 30 days. "
                 "Final rates and schedules are confirmed on the insertion order.", st_small)]
    doc.build(story)
    return buf.getvalue(), title


@app.get("/api/quotes/<int:qid>/pdf")
def quote_pdf(qid):
    db = SessionLocal()
    try:
        q = db.get(Quote, qid)
        if not q:
            return jsonify({"ok": False, "error": "Quote not found"}), 404
        state = json.loads(q.data or "{}")
        ensure_sections(state)
        pdf_bytes, title = build_proposal_pdf(q, state)
        q.pdf_blob = pdf_bytes
        q.pdf_filename = title + ".pdf"
        q.pdf_generated_at = datetime.now(timezone.utc)
        log_activity(db, q.id, "📄", f"PDF generated — {title}.pdf")
        db.commit()
        return send_file(BytesIO(pdf_bytes), mimetype="application/pdf",
                         as_attachment=request.args.get("download") == "1",
                         download_name=title + ".pdf")
    finally:
        db.close()


@app.get("/api/quotes/<int:qid>/pdf/archived")
def quote_pdf_archived(qid):
    """The last saved copy from the database (what 'save the PDF to the DB' returns)."""
    db = SessionLocal()
    try:
        q = db.get(Quote, qid)
        if not q or not q.pdf_blob:
            return jsonify({"ok": False, "error": "No archived PDF for this quote yet"}), 404
        return send_file(BytesIO(q.pdf_blob), mimetype="application/pdf",
                         as_attachment=False, download_name=q.pdf_filename or "proposal.pdf")
    finally:
        db.close()


# =====================================================================
# Word (.docx) export
# =====================================================================
def build_proposal_docx(q, state, sent_at=_UNSET):
    d = Document()
    for section in d.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    def hcolor(p, size, color, bold=True, center=False):
        run = p.runs[0] if p.runs else p.add_run("")
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = color
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    NAVY_D = RGBColor(0x14, 0x28, 0x4B)
    MUTED_D = RGBColor(0x53, 0x65, 0x7A)

    p = d.add_paragraph("SMART 1 MARKETING")
    hcolor(p, 16, NAVY_D, center=True)
    p = d.add_paragraph(f"Marketing Proposal — {q.client or 'Client'}")
    hcolor(p, 20, NAVY_D, center=True)
    p = d.add_paragraph(f"Quote {q.quote_number}  ·  {datetime.now().strftime('%B %d, %Y')}")
    hcolor(p, 10, MUTED_D, bold=False, center=True)

    sel = state.get("selectedPackage") or {}
    areas = campaign_areas(state)
    table = d.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for label, val in [("Client", q.client), ("Website", q.website),
                       ("Campaign Goals", "\n".join(state.get("objectives") or [])),
                       ("Target Area" if len(areas) < 2 else f"Target Areas ({len(areas)})",
                        "\n".join(hub_areas.names(areas)) or q.geo_summary),
                       ("Term", f"{q.months} months"),
                       ("Monthly campaign investment", _money(q.monthly_budget)),
                       (f"Total campaign investment ({q.months} months)",
                        _money(q.total_budget))]:
        # An empty row is dropped, the same reading the PDF's cover applies:
        # a labelled blank is a question the document forgot to answer.
        if not str(val or "").strip():
            continue
        row = table.add_row().cells
        row[0].text = label
        row[1].text = str(val or "")
        row[0].paragraphs[0].runs[0].font.bold = True

    sec_no = 0
    for sec in state.get("sections") or []:
        if not sec.get("enabled", True):
            continue
        # Numbered the way the PDF numbers them -- over the enabled sections,
        # at build time -- so the Word copy and the PDF of one proposal do
        # not disagree about which section is 07.
        sec_no += 1
        h = d.add_heading(f"{sec_no:02d}  {sec.get('title') or ''}", level=2)
        for run in h.runs:
            run.font.color.rgb = NAVY_D
        if sec.get("body"):
            _docx_body(d, sec.get("body"))
        kind = sec.get("kind")
        if kind == "cover":
            # The same standing trends block the PDF carries on the cover.
            t = hub_spec.industry_trends(state.get("industry"))
            _docx_body(d, "What is happening in "
                       + (t["industry"] if t["matched"] else "digital marketing")
                       + "\n" + "\n".join("• " + line for line in t["trends"])
                       + "\nHow Smart 1 helps\n" + t["help"]
                       + "\nHow budgets like this are usually crafted\n"
                       + t["budget"] + "\n" + t["note"])
        if kind == "areas":
            png, map_meta = campaign_map(state)
            if png:
                _docx_map(d, png, map_meta)
        table_plan = section_table(sec)
        if not table_plan["show"]:
            continue
        if table_plan["rows"] is not None:
            head = table_plan["head"] or []
            body = table_plan["rows"]
            width = max([len(r) for r in body] + ([len(head)] if head else []), default=0)
            if width:
                t = d.add_table(rows=1 if head else 0, cols=width)
                t.style = "Light Grid Accent 1"
                if head:
                    cells = t.rows[0].cells
                    for i, value in enumerate(head[:width]):
                        cells[i].text = value
                        if cells[i].paragraphs[0].runs:
                            cells[i].paragraphs[0].runs[0].font.bold = True
                for row in body:
                    cells = t.add_row().cells
                    for i, value in enumerate(row[:width]):
                        cells[i].text = value
            continue
        if kind == "areas" and areas:
            t0 = d.add_table(rows=1, cols=3)
            t0.style = "Light Grid Accent 1"
            hdr = t0.rows[0].cells
            for i, htxt in enumerate(["Target area", "Coverage", "Est. population"]):
                hdr[i].text = htxt
                hdr[i].paragraphs[0].runs[0].font.bold = True
            for area in areas:
                population = hub_areas.estimated_population(area)
                zips = hub_areas.zip_list(area.get("zips"))
                row = t0.add_row().cells
                row[0].text = hub_areas.label(area)
                row[1].text = area.get("notes") or (f"{len(zips)} ZIP Codes" if zips else area["type"])
                row[2].text = f"{population:,}" if population else "Not measured"
            for row in targets_of_interest(state):
                d.add_paragraph(
                    f"{row['name']} — "
                    + (row["note"] or ("their customers" if row["kind"] == "competitor"
                                       else "the people who go there"))
                    + (f" (geo-fenced at {row['address']})" if row["fenceable"]
                       else " (brand and behavior — no address on file to fence)"))
        elif kind == "mediaplan" and state.get("items"):
            # Same reading as the PDF and the preview. This table used to be
            # four columns to the PDF's five -- one client's proposal saying
            # two different things depending on which file was sent.
            plan = media_plan_rows(state)
            t2 = d.add_table(rows=1, cols=len(plan["columns"]))
            t2.style = "Light Grid Accent 1"
            hdr = t2.rows[0].cells
            heads = [*plan["columns"][:4], f"Total ({plan['months']} mo)",
                     plan["columns"][5]]
            for i, htxt in enumerate(heads):
                hdr[i].text = htxt
                hdr[i].paragraphs[0].runs[0].font.bold = True
            for r in plan["rows"]:
                row = t2.add_row().cells
                row[0].text = r["product"]
                row[1].text = r["category"]
                row[2].text = r["rate"]
                row[3].text = r["monthly_label"] or _money(r["monthly"])
                row[4].text = _money(r["campaign"])
                row[5].text = r["delivery"]
            total_row = t2.add_row().cells
            total_row[0].text = "Campaign total"
            total_row[3].text = _money(plan["monthly_total"])
            total_row[4].text = _money(plan["campaign_total"])
            if plan["note"]:
                d.add_paragraph(plan["note"])
        elif kind == "packages" and state.get("packages"):
            pkgs = state.get("packages") or []
            t3 = d.add_table(rows=1, cols=1 + len(pkgs))
            t3.style = "Light Grid Accent 1"
            hdr = t3.rows[0].cells
            hdr[0].text = ""
            for i, pkg in enumerate(pkgs):
                hdr[i + 1].text = pkg.get("name", "")
                hdr[i + 1].paragraphs[0].runs[0].font.bold = True
            for label, key in [("Monthly", "monthly"), ("Total", "total")]:
                row = t3.add_row().cells
                row[0].text = label
                for i, pkg in enumerate(pkgs):
                    row[i + 1].text = _money(pkg.get(key))
            # The same line the PDF prints under the same table: one client
            # note, so a Word copy and a PDF copy of one proposal cannot
            # disagree about how long the price stands.
            validity = hub_validity.client_note(_quote_window(q, state, sent_at))
            if validity:
                d.add_paragraph(validity)
        elif kind == "kpis" and state.get("kpis"):
            _docx_body(d, hub_spec.bullets(state.get("kpis") or []))
        elif kind == "friction":
            for pick in hub_discovery.suggestions(state):
                d.add_paragraph(f"We suggest they should {pick['title'][0].lower()}"
                                f"{pick['title'][1:]} — {pick['detail']}")
        elif kind == "channels":
            for item in channel_lines(state):
                role, kpi = _channel_role(item)
                d.add_paragraph(f"{item.get('product') or ''} — {role}. Measured on: {kpi}.")
        elif kind == "creative":
            for row in hub_creative.evaluate(state)["media"]:
                d.add_paragraph(f"{row['label']} — {_money(row['spend'])} campaign. "
                                f"{_creative_phrase(row)}. "
                                f"Needs: {hub_creative.units_line(state, row['medium'])}")
        elif kind == "roi":
            plan = hub_kpi.framework(state)
            if plan["measured"]:
                _docx_body(d, hub_spec.bullets([plan["primary"]], "Primary KPI"))
                if plan["secondary"]:
                    _docx_body(d, hub_spec.bullets(plan["secondary"], "Secondary KPIs"))
            else:
                d.add_paragraph(plan["note"])
            if plan["additional_metrics"]:
                _docx_body(d, hub_spec.bullets(plan["additional_metrics"],
                                               "Also tracked and reported monthly"))
            if plan["rows"]:
                kt = d.add_table(rows=1, cols=3)
                kt.style = "Light Grid Accent 1"
                head = kt.rows[0].cells
                for i, htxt in enumerate(["Product", "Key KPI", "Expected benchmark"]):
                    head[i].text = htxt
                    head[i].paragraphs[0].runs[0].font.bold = True
                for row in plan["rows"]:
                    cells = kt.add_row().cells
                    cells[0].text = row["product"]
                    cells[1].text = row["kpi"]
                    cells[2].text = row["expected"]
                d.add_paragraph("Benchmark ranges are what this inventory normally "
                                "delivers for campaigns of this shape. They are "
                                "expectations to measure against, not guarantees.")
        elif kind == "reach" and state.get("estimates"):
            est = state.get("estimates") or {}
            d.add_paragraph(f"Estimated population {int(est.get('pop') or 0):,} · addressable audience "
                            f"{int(est.get('aud') or 0):,} · households {int(est.get('hh') or 0):,} · "
                            f"devices {int(est.get('dev') or 0):,}")
        elif kind == "growth":
            growth = growth_options(state)
            if not growth["any"]:
                d.add_paragraph("Every product the discovery answers pointed at "
                                "is already on the plan above.")
            else:
                for row in growth["increases"]:
                    d.add_paragraph(
                        f"{row['product']}: quoted {_money(row['current'])}, "
                        f"recommended {_money(row['suggested'])} "
                        f"(+{_money(row['uplift'])}/mo, "
                        f"{_money(row['campaign'])} over {growth['months']} months)")
                for row in growth["additions"]:
                    price = (f"from {_money(row['suggested'])}/mo"
                             if row["quoted"] else "quoted on request")
                    d.add_paragraph(f"{row['product']} — {row['why']} ({price})")
        elif kind == "zips":
            wrote = False
            for area in campaign_areas(state):
                result = hub_areas.area_zips(area)
                if result["kept"]:
                    d.add_paragraph(f"{hub_areas.label(area)} — "
                                    f"{', '.join(result['kept'])}")
                    wrote = True
            if not wrote:
                d.add_paragraph("No ZIP Codes were captured for this campaign. "
                                "The insertion order needs them before trafficking.")
            for rule in hub_areas.zip_exceptions(state.get("targetAreas")
                                                 or campaign_areas(state)):
                d.add_paragraph(
                    f"{rule['area']} — exception: “{rule['text']}”. "
                    + (f"{rule['dropped']} ZIP Code(s) excluded, {rule['kept']} running."
                       if rule["applied"] else
                       "This could not be read and has NOT been applied; every ZIP above is running."))

    p = d.add_paragraph("Smart 1 Marketing · smart1marketing.com · This proposal is valid for 30 days.")
    hcolor(p, 8, MUTED_D, bold=False, center=True)
    buf = BytesIO()
    d.save(buf)
    return buf.getvalue()


@app.get("/api/quotes/<int:qid>/docx")
def quote_docx(qid):
    db = SessionLocal()
    try:
        q = db.get(Quote, qid)
        if not q:
            return jsonify({"ok": False, "error": "Quote not found"}), 404
        state = json.loads(q.data or "{}")
        ensure_sections(state)
        blob = build_proposal_docx(q, state)
        log_activity(db, q.id, "📝", f"Word copy exported — {q.quote_number}")
        db.commit()
        return send_file(BytesIO(blob),
                         mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                         as_attachment=True,
                         download_name=f"S1M Proposal - {q.quote_number} - {q.client or 'Client'}.docx")
    finally:
        db.close()


# =====================================================================
# Default proposal sections (used when the builder hasn't customized them)
# =====================================================================
# Sections the AI may write copy for. The table-backed ones take an intro
# paragraph above their generated table; only the cover has nothing to say.
WRITABLE_KINDS = ("text", "friction", "areas", "reach", "channels", "mediaplan",
                  "creative", "packages", "kpis", "roi", "growth")


def writable_sections(state):
    return [sec for sec in (state or {}).get("sections") or []
            if sec.get("kind") in WRITABLE_KINDS and sec.get("enabled", True)]


def industry_template(state):
    """The `hub.industries` entry that best fits this quote, or None.

    The library is keyed by our own industry slugs (boat, ski, legal...) while
    the wizard offers a broader list, so this matches on the label as well.
    Returning None is a normal answer — most industries have no template, and
    a generic proposal is better than one that claims a ski resort's demand
    triggers because the name nearly matched.
    """
    raw = str((state or {}).get("industry") or "").strip().lower()
    if not raw:
        return None
    if raw in hub_industries.INDUSTRIES:
        return hub_industries.INDUSTRIES[raw]
    for key, entry in hub_industries.INDUSTRIES.items():
        label = str(entry.get("label") or "").lower()
        # The key is matched on WORD boundaries, never as a substring: "rv"
        # is a substring of "services", so every Home Services and Financial
        # Services proposal opened its Executive Summary with the RV Dealers
        # pitch — "marketing that follows the camping calendar" on an HVAC
        # company's document, found by rendering one. hub/client_key.py's
        # rule, wearing an industry label.
        if raw == label or (len(raw) > 3 and
                            (raw in label
                             or re.search(rf"\b{re.escape(key)}\b", raw))):
            return entry
    return None


# =====================================================================
# Expected Results & ROI — computed, never written
# =====================================================================
def expected_results(state):
    """What the money buys, product by product, at the rates the client is quoted.

    No longer printed on the client's document: "Expected Results & ROI" is
    the KPI framework now, because an impression count says what the money
    bought rather than what the business gets, and the insertion order was
    already describing the campaign the second way. This stays because it is
    the one place that knows the delivery arithmetic — the quoted rate rather
    than the listed one, a one-time line spread across the flight rather than
    charged monthly, and a management fee reporting no units at all — and
    that arithmetic is what any future delivery figure has to come from.

    Directive: every proposal ends with this section. It is calculated here
    rather than asked of the model, because a projection written next to a
    media plan it contradicts is worse than no projection — and a model given
    a budget will produce impression counts that look authoritative and are
    invented.

    Anything without a CPM or CPV — a management fee, a flat monthly, a custom
    quote — reports no units at all rather than a plausible number. That is
    most of the value of doing it this way.
    """
    state = state or {}
    months = max(1, int(state.get("months") or 1))
    rows, totals = [], {"impressions": 0, "views": 0, "monthly": 0.0}
    unpriced = []

    for item in state.get("items") or []:
        try:
            monthly = float(item.get("dollars") or 0)
        except (TypeError, ValueError):
            monthly = 0.0
        # Monthly or one-time, and for how long. A line quoted at $3,000 that
        # is actually a one-off is $18,000 over a six-month flight if nobody
        # asks -- so the wizard asks, and the totals here read the answer
        # rather than assuming every line runs the whole campaign.
        basis = str(item.get("basis") or "monthly")
        try:
            term = int(item.get("termMonths") or months)
        except (TypeError, ValueError):
            term = months
        term = max(1, min(months, term))
        if basis == "one_time":
            line_campaign = round(monthly, 2)
            # A one-time cost still has to sit in a monthly plan somewhere, so
            # it is spread across the flight. Never dropped (the plan
            # under-reports what is owed) and never charged monthly (every
            # month of the campaign is overstated by the same amount).
            line_monthly = round(monthly / months, 2)
        else:
            line_campaign = round(monthly * term, 2)
            line_monthly = monthly
        totals["monthly"] += line_monthly
        product = hub_rate_card.find(item.get("label") or item.get("product") or "",
                                     item.get("category") or "")
        if product is None:
            # Fall back to the rate the wizard carried, so an off-card product
            # a rep added by hand is still estimated rather than dropped.
            product = {"rate_type": item.get("rate"), "rate_value": item.get("rateValue"),
                       "listed_rate": item.get("rate") or ""}
        # Delivery is what the client's money buys at the rate the client is
        # quoted. Estimated off the card's own rate it said a $1,000 line
        # buys 235,000 impressions at $4.25 -- printed on the client's ROI
        # table, and undeliverable, because the line is sold at $8.50 and
        # buys half of that.
        quoted = _sell_rate(item)
        priced = dict(product)
        if quoted:
            # The line's own rate wins over the looked-up row. `find()`
            # matches on the product name, and four categories carry a
            # product called "Demographic" -- so a location-lookback line
            # resolved to the $4.25 display row and its delivery was
            # estimated against a rate nobody quoted. The wizard carried the
            # rate off the card row the rep actually picked; that is the
            # one to price against.
            priced["rate_type"] = item.get("rate") or priced.get("rate_type")
            priced["rate_value"] = quoted
            priced["listed_rate"] = f"{item.get('rate')} ${quoted:,.2f}"
        delivery = hub_rate_card.estimate_delivery(priced, monthly)
        if quoted:
            delivery["note"] = (f"At the quoted rate of {item.get('rate')} "
                                f"${quoted:,.2f}.")
        units = delivery.get("units")
        run = 1 if basis == "one_time" else term
        if units and delivery["unit_label"].startswith("impressions"):
            totals["impressions"] += units * run
        elif units:
            totals["views"] += units * run
        if not units:
            unpriced.append(hub_rate_card.quote_label(item.get("product"),
                                                      item.get("category")))
        rows.append({
            "product": hub_rate_card.quote_label(item.get("product"),
                                                 item.get("category")),
            "category": item.get("category") or "",
            "medium": hub_creative.medium_of(item),
            "quoted_rate": quoted,
            "monthly": line_monthly,
            "campaign": line_campaign,
            "basis": basis,
            "term_months": 1 if basis == "one_time" else term,
            "rate": (f"{item.get('rate')} ${quoted:,.2f}" if quoted
                     else (product.get("listed_rate") or "")),
            "units": units,
            "unit_label": delivery.get("unit_label") or "",
            "note": delivery.get("note") or "",
        })

    totals["campaign"] = round(sum(r["campaign"] for r in rows), 2)
    return {
        "rows": rows, "months": months, "totals": totals,
        # Only says anything when they run traditional media and a posture has
        # been chosen; otherwise "" and the section renders without it.
        "traditional_note": hub_discovery.roi_note(state),
        # Named, so the section can say which products are not represented in
        # the headline number instead of quietly under-reporting.
        "unpriced": unpriced,
        "metrics": _tracked_metrics(state),
    }


# Column headings for the media plan, in one place because three renderers
# draw this table and they had already drifted: the PDF carried Rate and the
# Word export did not, so the same client's proposal said two different things
# depending on which file was sent.
MEDIA_PLAN_COLUMNS = ("Product", "Category", "Rate", "Monthly", "Total",
                      "Delivery / mo")


def campaign_cost(state) -> dict:
    """What this campaign costs, once, for every screen that prints a number.

    ## Why this exists

    One proposal used to carry three different monthly figures and hand a
    fourth to the insertion order. `summarize_into()` took `monthly_budget`
    from the selected package or from `state["budget"]` -- the number the rep
    typed on the Budget step, which is what the *client asked for* -- while
    the media plan table totalled the lines that are actually being bought
    and `ioDataPayload()` billed those same lines. So a plan edited down from
    $8,000 to $5,750 produced a cover reading "$8,000 / mo", a media mix
    totalling $5,750, an investment summary quoting $8,000 of media, and an
    insertion order for $5,750: a client signing a document $2,250 a month
    away from the order that bills them, with every screen internally
    consistent and nothing erroring anywhere.

    **The plan is the number.** Once there are line items they are what is
    being bought, and every other figure is derived from them.

    ## The rules

      * **Recurring and one-time are never added together.** A $1,500 video
        shoot is not $1,500 a month, and folding it into the monthly figure
        overstates every month of the campaign -- which is exactly what the
        old `sum(i["dollars"])` fallback did.
      * **A line runs for its own term.** `expected_results()` already reads
        `termMonths`; the same reading is used here rather than a second one.
      * **The Suite licence is not campaign cost.** It is a separate product
        with its own line in the investment summary, and blending it is how a
        client comes to believe the platform stops costing money when they
        pause the media -- the rule `hub/proposal_spec.py` states.
      * **A quote with no plan yet still answers.** `stated` carries what the
        client asked for, so a proposal at the Budget step reads sensibly and
        nothing has to branch on "are there items".
    """
    state = state or {}
    months = max(1, int(state.get("months") or 1))
    recurring = one_time = campaign = 0.0
    monthly_lines, one_time_lines = [], []
    for item in state.get("items") or []:
        try:
            dollars = float(item.get("dollars") or 0)
        except (TypeError, ValueError):
            dollars = 0.0
        label = hub_rate_card.quote_label(item.get("product"),
                                          item.get("category"))
        if str(item.get("basis") or "monthly") == "one_time":
            one_time += dollars
            campaign += dollars
            one_time_lines.append({"label": label, "amount": round(dollars, 2)})
            continue
        try:
            term = int(item.get("termMonths") or months)
        except (TypeError, ValueError):
            term = months
        term = max(1, min(months, term))
        recurring += dollars
        campaign += dollars * term
        monthly_lines.append({"label": label, "amount": round(dollars, 2),
                              "term_months": term})

    try:
        stated = float(state.get("budgetAsked") or state.get("budget") or 0)
    except (TypeError, ValueError):
        stated = 0.0
    has_plan = bool(monthly_lines or one_time_lines)
    return {
        "months": months,
        "has_plan": has_plan,
        # The recurring campaign cost: media and the services bought beside
        # it, which is what a rep means by "the campaign" and what the
        # insertion order bills every month.
        "recurring": round(recurring, 2),
        "one_time": round(one_time, 2),
        "campaign": round(campaign, 2),
        "monthly_lines": monthly_lines,
        "one_time_lines": one_time_lines,
        # What the client asked for, kept whatever the plan became. Losing it
        # would lose the only record of the conversation the tiers were sized
        # against.
        "stated": round(stated, 2),
        "differs_from_stated": bool(has_plan and stated
                                    and abs(stated - recurring) >= 1),
    }


_FLAT_RATE_RE = re.compile(
    r"^\$?\s*(\d[\d,]*(?:\.\d+)?)\s*(?:flat|one[- ]time|/mo|per month|monthly)?$",
    re.IGNORECASE)


def _plan_rate(row) -> str:
    """What the Budget Allocation table's rate column says for one line.

    A quoted CPM/CPV keeps its rate ("CPM $8.50") and a flat fee keeps its
    figure, formatted as money. Everything else is "Managed": a percentage,
    a custom quote or a sentence about the card is our pricing described,
    not a rate stated, and on this table it reads as a figure the client
    should be able to check and cannot.
    """
    if row.get("quoted_rate"):
        return row.get("rate") or "Managed"
    m = _FLAT_RATE_RE.match(str(row.get("rate") or "").strip())
    if m:
        try:
            return f"${float(m.group(1).replace(',', '')):,.2f} flat"
        except ValueError:
            pass
    return "Managed"


def media_plan_rows(state) -> dict:
    """The media plan as every renderer draws it, computed once.

    **Delivery is back on the client's document, and this is where it belongs.**
    It came off the proposal with "Expected Results & ROI", correctly: an
    impression count answers what the money *bought*, not what the business
    gets, and printed under that heading it read as a promise about outcomes.
    Under the media plan it is answering the question the media plan asks --
    what does this line buy -- and a client comparing two proposals has no
    other way to tell a $4.25 CPM apart from an $8.50 one.

    It is `expected_results()`'s arithmetic and not a second copy: the quoted
    rate rather than the listed one, a one-time line spread across the flight,
    and a management fee reporting **no units at all** rather than a plausible
    number. That last one is most of the value of doing it this way.
    """
    est = expected_results(state)
    months = est["months"]
    rows = []
    for r in est["rows"]:
        units, label = r.get("units"), r.get("unit_label") or ""
        unit = label.replace("/month", "").strip()
        # The rate column carries a rate a client can do arithmetic with -- a
        # quoted CPM/CPV or a flat fee -- and says "Managed" for everything
        # else. The card's own strings for the rest ("15% of gross", "Custom
        # quote", a sentence about the rate card) describe our pricing rather
        # than state a rate, and on the one table headed Budget Allocation
        # they read as figures a client should be able to check and cannot.
        # A bare flat number is formatted on the way through, because "250.0"
        # printed at a client is the bare-float failure one release back.
        rate = _plan_rate(r)
        if not units:
            # Never a zero and never a dash on its own: a fee that buys no
            # impressions is a different statement from one nobody priced.
            delivery = "Not impression-based"
        elif (r.get("basis") or "monthly") == "one_time":
            # A one-time line does not deliver every month, and printing its
            # units under a per-month heading multiplies it by the flight.
            delivery = f"{units:,} {unit}, once"
        else:
            # The column's own heading carries "per month"; repeating it on
            # every row makes the narrowest column the widest.
            delivery = f"{units:,} {unit}"
        one_time = (r.get("basis") or "monthly") == "one_time"
        rows.append({
            "product": r["product"], "category": r["category"],
            "rate": rate,
            # `expected_results()` spreads a one-time cost across the flight,
            # which is right for its own arithmetic and wrong in a column
            # headed Monthly: a $1,500 shoot shown as $250 a month made the
            # media plan's monthly total $5,750 where the investment summary
            # said $5,500 recurring plus $1,500 once. The row says which it
            # is and the Monthly column carries only what recurs.
            "monthly": None if one_time else r["monthly"],
            "monthly_label": "One-time" if one_time else None,
            "campaign": r["campaign"],
            "units": units, "unit_label": label, "delivery": delivery,
            "basis": r.get("basis") or "monthly",
        })
    # The totals are campaign_cost()'s, so this table, the cover, the
    # investment summary, the insertion order and the dashboard all print one
    # number for what the campaign costs.
    cost = campaign_cost(state)
    return {
        "columns": list(MEDIA_PLAN_COLUMNS),
        "months": months,
        "rows": rows,
        "monthly_total": cost["recurring"],
        "campaign_total": cost["campaign"],
        "one_time_total": cost["one_time"],
        "impressions": (est["totals"].get("impressions") or 0),
        "views": (est["totals"].get("views") or 0),
        # Named, so the table can say which lines are not in the delivery
        # column rather than quietly under-reporting the campaign.
        "unpriced": est.get("unpriced") or [],
        "note": delivery_note(est),
    }


def delivery_note(est) -> str:
    """The sentence under the delivery column. An estimate printed with no
    words around it reads as a guarantee, which is the failure the ROI section
    was rebuilt to undo -- so the words travel with the figures."""
    est = est or {}
    totals = est.get("totals") or {}
    bits = []
    if totals.get("impressions"):
        bits.append(f"{int(totals['impressions']):,} impressions")
    if totals.get("views"):
        bits.append(f"{int(totals['views']):,} views")
    if not bits:
        return ("No line in this plan is bought on an impression or view rate, "
                "so there is no delivery estimate to show.")
    note = ("Delivery is estimated at the rates quoted above and is what this "
            "plan is expected to buy over the campaign — "
            + " and ".join(bits) + " — not a guarantee of results.")
    unpriced = [u for u in (est.get("unpriced") or []) if u]
    if unpriced:
        one = len(unpriced) == 1
        note += (" " + ", ".join(unpriced)
                 + (" is not bought on an impression rate, so it is"
                    if one else
                    " are not bought on impression rates, so they are")
                 + " not counted in that figure.")
    return note


def _tracked_metrics(state):
    """What we will report on — the shared reading, not a second copy.

    This was its own list until the ROI section became the KPI framework, and
    two answers to "what do we report on" is how the proposal and the
    insertion order came to describe one campaign differently in the first
    place.
    """
    return hub_kpi.success_metrics(state)


# =====================================================================
# Proposal sections — the Smart 1 13-part structure
# =====================================================================
def ensure_sections(state):
    """The proposal's sections, seeded from the Smart 1 specification.

    `hub.proposal_spec.OUTLINE` owns the structure; this fills the prose the
    rep can then edit or have the AI write. A quote saved against the old
    eight-section layout keeps its sections — its copy is real work — but
    gains any *required* section it is missing, because a proposal without
    Expected Results & ROI is one the specification does not permit.
    """
    existing = state.get("sections")
    if existing:
        # A retired section comes OFF a stored quote here, whatever release
        # saved it -- this runs on every save and in front of the PDF and the
        # Word export, so a quote nobody re-opens still stops printing it.
        # Matched on the kind, because the id is editable and the kind is
        # what the renderers dispatch on.
        retired = [sec for sec in existing if isinstance(sec, dict)
                   and str(sec.get("kind")) in hub_spec.RETIRED_SECTION_KINDS]
        if retired:
            existing = [sec for sec in existing if sec not in retired]
            state["sections"] = existing
        have = {str(sec.get("id")) for sec in existing if isinstance(sec, dict)}
        seeded = _seeded_sections(state)
        for spec in hub_spec.OUTLINE:
            if spec["id"] in hub_spec.REQUIRED and spec["id"] not in have:
                addition = next(s for s in seeded if s["id"] == spec["id"])
                existing.append(addition)
        return state
    state["sections"] = _seeded_sections(state)
    return state


def _seeded_sections(state):
    client = state.get("client") or "the client"
    goals = ", ".join(state.get("objectives") or []) or "the campaign goals"
    areas = campaign_areas(state)
    where = hub_areas.summary(areas, limit=3) or "the target area"
    template = industry_template(state)
    months = max(1, int(state.get("months") or 1))
    segments = hub_spec.audience_segments_for(state.get("industry", ""))

    sections = hub_spec.default_sections()
    body = {}

    body["summary"] = (
        f"{client} has the demand; what is missing is a single system that captures it "
        f"and proves what it produced. This plan puts Smart 1 media in front of the "
        f"right households in {where} and routes every response into the Smart 1 Suite, "
        f"where it is answered, nurtured and measured. Over {months} months the goal is "
        f"{goals.lower()} — reported against business outcomes, not impressions.")
    if template:
        body["summary"] = template["intro"] + "\n\n" + body["summary"]

    body["objectives"] = (
        "Primary\n" + "\n".join(f"• {o}" for o in (state.get("objectives") or
                                                    ["To be confirmed with the client"]))
        + "\n\nSecondary\n"
        "• A lower, measurable cost per acquisition\n"
        "• One vendor and one dashboard instead of a stack of logins\n"
        "• Reporting the client can open themselves, at any time")

    body["friction"] = (
        "Most of the money lost in local marketing is lost after the click, not before "
        "it. The patterns worth checking here:\n"
        "• A bolted-on tech stack — separate tools for email, forms, reviews and calls, "
        "each with its own login and none of them talking to each other.\n"
        "• No central CRM, so an inbound lead cools while it waits for someone to "
        "notice it.\n"
        "• Fixed-schedule advertising that runs the same way regardless of weather, "
        "foot traffic or real-time intent.")

    if segments:
        body["areas"] = (
            "Targeting is built area by area from named third-party segments rather "
            "than broad demographics:\n"
            + "\n".join(f"• {seg}" for seg in segments))

    if template and template.get("channels"):
        body["channels"] = (
            "Each channel below has a job in the customer journey:\n"
            + "\n".join(f"• {ch}" for ch in template["channels"]))

    if template and template.get("triggers"):
        body["channels"] = (body.get("channels", "") +
                            "\n\nBudget is concentrated on the moments that produce "
                            "revenue: " + ", ".join(template["triggers"]) + ".").strip()

    # Deliberately says nothing about markup. This read "every rate is the
    # Smart 1 card rate -- there is no markup between the line item and what
    # runs", which named our internal pricing on a document a client reads
    # and, since sell_rate() started quoting CPM at 2x, was printed directly
    # above a table disproving it.
    body["mediaplan"] = (
        "The split below is weighted toward the stage of the funnel this campaign has "
        "to move, and the delivery column is what each line is expected to buy at the "
        "rate it is quoted at.")

    body["creative"] = _creative_section_body(state)

    body["technology"] = (
        "The Smart 1 Suite is the central nervous system of this campaign. Every call, "
        "form, chat and message the media generates lands in one inbox, and the Suite "
        "goes to work immediately: Missed Call Text Back so a missed call becomes a "
        "conversation instead of a lost lead, automated text and email follow-up, "
        "online scheduling, and automated review requests that compound into local "
        "search visibility. The media creates the opportunity; the Suite is what turns "
        "it into revenue you can trace.")

    body["reporting"] = (
        "Optimization is a routine, not a promise: bids adjusted against delivery and "
        "cost per action, negative keywords and placement exclusions updated, creative "
        "checked against performance and rotated before it fatigues. Everything reports "
        "into one live dashboard inside the Smart 1 Suite that you can open whenever "
        "you want, rather than waiting for a monthly PDF.")

    body["packages"] = (
        "Platform licensing, media spend and any one-time production are listed "
        "separately below so it is clear what recurs and what does not.")

    body["roi"] = (
        "This is what the campaign will be judged on, product by product, with the "
        "result each one normally delivers. The Smart 1 Suite is the single source "
        "of truth for what the media produced: every lead is attributed to the "
        "channel that created it, so the spend can be judged against the business "
        "rather than against a click count.")

    body["next"] = "\n".join(f"{i}. {step}" for i, step in
                              enumerate(hub_spec.NEXT_STEPS, 1))

    for section in sections:
        if section["id"] in body and body[section["id"]]:
            section["body"] = body[section["id"]]
    return sections


def _creative_section_body(state):
    """Messaging themes plus what the creative gate actually decided."""
    lines = [
        "Three messaging themes carry the campaign, rotated so no one execution "
        "fatigues:",
        "• Problem / Agitate / Solution — name the frustration the customer already "
        "has, make it concrete, then resolve it.",
        "• Proof — the specific, verifiable reason to choose this business over the "
        "one down the road.",
        "• Urgency without gimmick — the real reason now is better than later "
        "(season, capacity, availability), never a manufactured countdown.",
    ]
    summary = hub_creative.summary_line(state)
    if summary:
        lines += ["", "Creative source: " + summary + "."]
    return "\n".join(lines)


# =====================================================================
# AI routes (optional — need OPENAI_API_KEY; same pattern as the IO app)
# =====================================================================
from hub import openai_responses as _responses


def _openai_call(payload, api_key):
    """The transport, kept as this module's own name so a test can stand in
    front of it. The reading itself is `hub/openai_responses.py` — the IO
    Builder carried a second copy of it, and the copy was the bug: the hosted
    web-search tool rode on every call there too, and the fix that landed here
    never reached it."""
    return _responses.post(payload, api_key)


def _openai_response(prompt, max_output_tokens=6000, search=False):
    """One call to the Responses API, with three ways of failing named.

    ``search`` is opt-in, a refusal of the hosted tool falls back without it,
    a refusal carries the API's own sentence and an answer cut short is said
    to be that. All four rules, and why each exists, are in
    `hub/openai_responses.py`.
    """
    return _responses.ask(prompt, module="sales_builder", purpose="quote",
                          max_output_tokens=max_output_tokens, search=search,
                          call=_openai_call)



def _json_from_ai(text):
    cleaned = text.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?\s*|\s*```$", "", cleaned, flags=re.I | re.S)
    return json.loads(cleaned)


@app.post("/api/ai/draft-proposal")
def ai_draft_proposal():
    """One sentence -> a pre-filled proposal draft the salesperson reviews."""
    body = request.get_json(force=True) or {}
    ask = str(body.get("prompt") or "").strip()
    if not ask:
        return jsonify({"ok": False, "error": "A prompt is required"}), 400
    catalog = body.get("categories") or []
    prompt = (
        "You are a senior media planner at Smart 1 Marketing. From the request below, draft a campaign "
        "proposal setup. Return STRICT JSON only with keys: client (string), url (string, may be empty), "
        "industry (one of: Healthcare / Dental, Home Services, Legal, Automotive, Real Estate, Retail / "
        "E-comm, Restaurant / Hospitality, Financial Services, Fitness / Wellness, Education, B2B / "
        "Professional, Other), objectives (array from: Brand Awareness, Lead Generation, Website Traffic, "
        "Store Visits, Phone Calls, Conversions, Recruitment, Event Promotion), geoType (one of: City/ZIP "
        "+ Radius, DMA, Statewide, National, Other), geo (string), radius (number, only for radius type), "
        "budget (number, monthly dollars), months (number), categories (array of up to 4 product "
        "categories chosen ONLY from this list: " + ", ".join(catalog) + "), rationale (short string).\n\n"
        "Request: " + ask)
    try:
        result = _json_from_ai(_openai_response(prompt, 3000))
        return jsonify({"ok": True, "draft": result})
    except Exception as exc:
        logger.exception("AI draft failed")
        return jsonify({"ok": False, "error": "AI draft failed", "detail": str(exc)}), 502


@app.post("/api/ai/rewrite")
def ai_rewrite():
    """Rewrite a proposal section's copy per an instruction."""
    body = request.get_json(force=True) or {}
    text = str(body.get("text") or "").strip()
    instruction = str(body.get("instruction") or "make it clearer and more client-friendly").strip()
    if not text:
        return jsonify({"ok": False, "error": "Text is required"}), 400
    # The standing directives apply to a rewrite exactly as they do to a first
    # draft. They were not applied here at all, so "make it punchier" could
    # walk copy straight past the Smart 1 Labs exclusion or invent a statistic.
    guidance = hub_spec.guidance_for(str(body.get("section") or ""))
    prompt = (hub_spec.system_prompt(body.get("data") or {}) + "\n\n"
              "Rewrite the proposal section below. Instruction: " + instruction + ".\n"
              + (f"What this section is for: {guidance}\n" if guidance else "")
              + "Return only the rewritten text, no preamble.\n\n" + text)
    try:
        rewritten = _openai_response(prompt, 2500)
    except Exception as exc:                            # noqa: BLE001
        return jsonify({"ok": False, "error": "AI rewrite failed", "detail": str(exc)}), 502
    rewritten = hub_spec.clean_ai_text(rewritten)
    problems = hub_spec.violations(rewritten)
    if problems:
        # Returned unchanged rather than silently handing back copy that
        # breaks a rule: the rep asked for a rewrite, not for a rule waiver.
        return jsonify({"ok": True, "text": text, "warnings": problems
                        + ["The rewrite was discarded and your text kept."]})
    return jsonify({"ok": True, "text": rewritten, "warnings": []})


@app.post("/api/ai/draft-sections")
def ai_draft_sections():
    """Write the proposal's narrative copy from the campaign already built.

    The standalone builder's one real advantage was this: it produced a
    proposal a rep could send, not a form a rep then had to write into. It did
    it from an industry template and nothing else, though, so the prose never
    matched the media plan underneath it. This writes from both — the industry
    library for voice and demand triggers, the actual products, budget, target
    areas and KPIs for substance.

    Only `text` sections are written. The tables (media plan, packages, reach,
    target areas) are generated from the campaign data and must never be
    narrated by a model — that is how a proposal ends up describing a budget
    it does not contain.
    """
    body = request.get_json(force=True) or {}
    state = body.get("data") or {}
    ensure_sections(state)
    template = industry_template(state)
    areas = campaign_areas(state)

    # One section per request when the caller names one. The wizard writes the
    # proposal a section at a time so its loader can say which section is
    # being written rather than spinning on "generating…" for a minute, and so
    # a single failed section does not cost the other twelve.
    only = str(body.get("section") or "").strip()
    writable = [sec for sec in state.get("sections") or []
                if sec.get("kind") in WRITABLE_KINDS and sec.get("enabled", True)
                and (not only or sec.get("id") == only)]
    if not writable:
        return jsonify({"ok": False, "error": "This proposal has no text sections to write."}), 400

    facts = {
        "client": state.get("client") or "",
        "website": state.get("url") or "",
        "industry": state.get("industry") or "",
        "business_description": state.get("description") or "",
        "objectives": state.get("objectives") or [],
        "kpis": state.get("kpis") or [],
        "target_areas": hub_areas.for_prompt(areas),
        "target_area_count": len(areas),
        "audiences": state.get("audiences") or [],
        "exclusions": state.get("exclusions") or [],
        "months": state.get("months") or 1,
        "monthly_budget": (state.get("selectedPackage") or {}).get("monthly") or state.get("budget") or 0,
        "products": [{"product": i.get("product"), "category": i.get("category"),
                      "monthly": i.get("dollars")} for i in (state.get("items") or [])],
        "landing_page": state.get("landingUrl") or "",
        "landing_recommendation": state.get("landing") or "",
        "creative_source": hub_creative.summary_line(state),
        # What they already run, what they don't, and where they stand on
        # their traditional media. `traditional_guidance` carries the posture
        # AND the instruction not to argue it -- a model handed "they want to
        # shift budget to digital" writes a case against radio, and a proposal
        # that opens by calling a client's existing spend wasted loses the
        # room before the media plan is read.
        "current_marketing": hub_discovery.for_prompt(state),
        "traditional_guidance": hub_discovery.guidance(state),
        "we_suggest": [s["title"] for s in hub_discovery.suggestions(state)],
        # Named competitors and places. Without these the audience section
        # says "competitor conquesting", which is a tactic rather than a
        # campaign, and is the paragraph a client checks hardest.
        "targets_of_interest": [
            {"name": row["name"], "what_we_want": row["note"],
             "reached_by": ("geo-fenced at their address" if row["fenceable"]
                            else "brand and behavior targeting — no address on file")}
            for row in targets_of_interest(state)],
        # What the Suite licence is on the quote for, in this client's own
        # terms. The technology section used to describe the Suite in the
        # abstract on every proposal ever produced here, which is why nobody
        # read it.
        "suite_covers": hub_discovery.suite_line(
            state, (state.get("suiteTier") or {}).get("name") or ""),
        "zip_exceptions": [r["note"] for r in
                           hub_areas.zip_exceptions(campaign_areas(state))
                           if r["applied"]],
        "suite_tier": hub_spec.suggested_tier(
            (state.get("selectedPackage") or {}).get("monthly")
            or state.get("budget") or 0)["name"],
        "sections_to_write": [{"id": sec.get("id"), "title": sec.get("title"),
                               "guidance": hub_spec.guidance_for(sec.get("id")),
                               "has_table": sec.get("kind") not in ("text", "cover"),
                               "current": sec.get("body") or ""} for sec in writable],
    }
    if template:
        facts["industry_voice"] = template["intro"]
        facts["demand_triggers"] = template["triggers"]
        facts["channels_we_sell_here"] = template["channels"]

    prompt = (
        # The standing directives, the audience segments and the operating
        # facts that apply to this campaign all come from hub.proposal_spec,
        # so the rules land in one place rather than in a prompt literal that
        # drifts every time someone edits this route.
        hub_spec.system_prompt(state) + "\n\n"
        "Write the narrative copy for the client-facing proposal described below.\n\n"
        "Output rules:\n"
        "- Return STRICT JSON only: {\"sections\": [{\"id\": \"...\", \"body\": \"...\"}]}, one entry per "
        "id in sections_to_write, and no other keys.\n"
        "- Each section answers the `guidance` given for it. Follow that guidance.\n"
        "- A section marked has_table:true is followed by a generated table of "
        "real numbers. Write the short paragraph that INTRODUCES it — say what "
        "the reader is about to look at and why it is shaped this way. Never "
        "restate the figures; they are directly below your copy and any number "
        "you write that differs from one of theirs is the error that matters.\n"
        "- Each section is 2 to 4 short paragraphs, plain professional English. No "
        "headings; keep bullet characters only where the current copy uses them.\n"
        "- Keep anything factual the current copy already states.\n\n"
        + json.dumps(facts, ensure_ascii=False))
    try:
        result = _json_from_ai(_openai_response(prompt, 2000 if only else 6000))
        written = {str(sec.get("id")): str(sec.get("body") or "")
                   for sec in (result.get("sections") or []) if sec.get("id")}
    except Exception as exc:                            # noqa: BLE001
        logger.exception("AI section draft failed")
        return jsonify({"ok": False, "error": "AI draft failed", "detail": str(exc)}), 502
    if not written:
        return jsonify({"ok": False, "error": "The AI returned no sections."}), 502

    # A directive is a rule, not a request. Copy that breaks one is dropped
    # rather than shown to a rep who would have to notice it -- the Smart 1
    # Labs exclusion is the whole reason this check exists.
    breaches = {}
    for sec_id, text in list(written.items()):
        # Markdown and emoji are cleaned rather than discarded. A section is
        # thrown away for breaking a standing directive -- naming Smart 1
        # Labs -- and a stray asterisk is not that: dropping otherwise good
        # copy over one would cost the rep the section and tell them nothing.
        written[sec_id] = text = hub_spec.clean_ai_text(text)
        problems = hub_spec.violations(text)
        if problems:
            breaches[sec_id] = problems
            del written[sec_id]
    if not written:
        return jsonify({"ok": False, "error": "Every section the AI returned broke a "
                                              "proposal rule and was discarded.",
                        "breaches": breaches}), 502
    return jsonify({"ok": True, "sections": written, "breaches": breaches})


# =====================================================================
# Campaign intelligence — served here, not fetched from the IO app
# =====================================================================
@app.post("/api/creative-check")
def api_creative_check():
    """What the creative gate still needs to be told about this media plan.

    Server-side as well as in the wizard, because the answer travels onto the
    insertion order and into the PDF -- and because the wizard's classifier is
    a mirror of `hub.creative_needs`, not the authority on it.
    """
    body = request.get_json(force=True) or {}
    state = body.get("data") or body
    result = hub_creative.evaluate(state)
    result["comp_confirm_under"] = hub_creative.COMP_CONFIRM_UNDER
    result["typical_production"] = hub_creative.TYPICAL_PRODUCTION
    result["summary"] = hub_creative.summary_line(state)
    return jsonify({"ok": True, **result})


@app.get("/api/proposal-spec")
def api_proposal_spec():
    """The Smart 1 proposal specification the wizard builds against."""
    return jsonify({
        "ok": True,
        "outline": [{"id": s["id"], "title": s["title"], "kind": s["kind"],
                     "purpose": s["purpose"], "required": s["id"] in hub_spec.REQUIRED}
                    for s in hub_spec.OUTLINE],
        "saas_tiers": hub_spec.SAAS_TIERS,
        "next_steps": hub_spec.NEXT_STEPS,
        "comp_confirm_under": hub_creative.COMP_CONFIRM_UNDER,
        "typical_production": hub_creative.TYPICAL_PRODUCTION,
        # The joint minimum rule, so the wizard blocks exactly what the IO
        # blocks rather than carrying a second opinion about the same number.
        "minimums": hub_rate_card.minimums_for_js(),
    })


@app.post("/api/ai/section-plan")
def api_section_plan():
    """The sections the wizard should ask for, in order, with their labels.

    Returned rather than derived in JavaScript so the loader counts exactly
    what the server will write -- a progress bar that says "3 of 9" while the
    server writes 13 is worse than no progress bar.
    """
    body = request.get_json(force=True) or {}
    state = body.get("data") or {}
    ensure_sections(state)
    return jsonify({"ok": True, "sections": [
        {"id": sec.get("id"), "title": sec.get("title"),
         "has_table": sec.get("kind") not in ("text", "cover"),
         "has_copy": bool((sec.get("body") or "").strip())}
        for sec in writable_sections(state)]})


@app.get("/api/industries")
def api_industries():
    """The industry library, for the wizard's picker."""
    return jsonify({"ok": True, "industries": hub_industries.industry_list()})


@app.post("/api/generate-business-description")
def api_business_description():
    """A Google Business Profile description, to the IO builder's rules.

    Identical prompt to /tools/io — it lives in hub.business_description now.
    Before this route existed, the wizard's Generate button called a path on
    the IO *app* that has never had it, so the placeholder ran every time and
    the description on a proposal was never the one on the IO.
    """
    body = request.get_json(force=True) or {}
    urls = [str(u).strip() for u in (body.get("urls") or []) if str(u).strip()]
    if not urls:
        return jsonify({"ok": False, "error": "A website URL is required."}), 400
    areas = hub_areas.normalize(body.get("areas") or body.get("targetAreas")) \
        or hub_areas.from_legacy(body)
    prompt = hub_desc.prompt_for(urls, client=body.get("client", ""),
                                 industry=body.get("industry", ""), areas=areas,
                                 brandfetch=body.get("brandfetch") or {},
                                 geo=str(body.get("geo") or ""))
    try:
        description = _openai_response(prompt, 5000)
    except Exception as exc:                            # noqa: BLE001
        return jsonify({"ok": False, "error": "Description request failed",
                        "detail": str(exc)}), 502
    # A Google Business Profile description is pasted straight into Google,
    # which renders no markup at all -- so this one is cleaned to plain text
    # with the bold markers taken out as well.
    description = hub_spec.plain_text(description)
    if not description:
        return jsonify({"ok": False, "error": "The AI returned no description."}), 502
    return jsonify({"ok": True, "description": description,
                    "warnings": hub_desc.check(description)})


def _headings_line(observed):
    """Read from `modules/ads_builder/landing_page.py`, which is where the
    shape it describes comes from. The IO Builder's landing review needs the
    same line, and a second copy of it drifts the day either end changes."""
    from modules.ads_builder.landing_page import headings_line
    return headings_line(observed)


@app.post("/api/review-landing-page")
def api_review_landing_page():
    """Conversion review of the landing page, before the campaign is priced.

    **The page is fetched.** It used to be a URL handed to a model with the
    word "Visit", which no model here can do -- so the answer was either a
    confident review of a page nobody had looked at, or, once the model was
    honest about it, the criteria it would have used followed by a sentence
    saying it could not reach the site. Both read to a rep as a broken
    button, and the first is worse than broken: it is fiction quoted to a
    client.

    ``modules.ads_builder.landing_page`` already does this properly for Smart
    1 Ads -- it requests the page and counts the conversion points off the
    markup, each carrying the evidence found -- so it is read here rather
    than copied. The observed half is fact and the model is given the facts
    and asked only for judgment, and the two are kept apart in the response
    so the screen can say which is which.
    """
    body = request.get_json(force=True) or {}
    url = str(body.get("url") or "").strip()
    if not url:
        return jsonify({"ok": False, "error": "A landing-page URL is required."}), 400

    try:
        from modules.ads_builder import landing_page as _lp
    except Exception as exc:                            # noqa: BLE001
        return jsonify({"ok": False, "error": "The page reader is unavailable",
                        "detail": str(exc)}), 502

    observed = _lp.observe(url)
    if not observed.get("measured"):
        # A page we could not fetch is reported as a page we could not fetch,
        # with the status behind it. Asking the model anyway is how a review
        # of a 404 gets written and pasted into a proposal.
        return jsonify({"ok": False, "url": url, "observed": observed,
                        "error": "That page could not be read.",
                        "detail": observed.get("error")
                                  or "The site did not answer."}), 502

    points = observed.get("conversion_points") or []
    facts = "\n".join(f"- {p['label']}: {p['evidence']}" for p in points[:25]) \
        or "- none found on the page"
    prompt = (
        "Review a campaign landing page as a conversion-focused page. Everything below was read "
        "off the live page just now -- treat it as fact, do not contradict it, and do not describe "
        "anything that is not in it.\n\n"
        f"URL: {observed.get('url') or url}\n"
        f"Client: {str(body.get('client') or '')}\n"
        f"Product or use: {str(body.get('product') or 'Campaign landing page')}\n"
        f"Campaign goals: {', '.join(str(o) for o in (body.get('objectives') or []))}\n"
        f"Page title: {observed.get('title') or '(none)'}\n"
        f"Meta description: {observed.get('meta_description') or '(none)'}\n"
        f"Declares a mobile viewport: "
        f"{'yes' if observed.get('mobile_viewport') else 'no' if observed.get('mobile_viewport') is False else 'not measured'}\n"
        f"Headings, in order: {_headings_line(observed) or '(none)'}\n"
        f"Conversion points found on the page:\n{facts}\n\n"
        f"Page text:\n{observed.get('text') or '(no readable text)'}\n\n"
        "Return a concise internal note with these headings: CTA Status, Strengths, Required Fixes "
        "Before Launch, Recommended Improvements, Tracking Checks. Be specific and practical, and "
        "quote what is actually on the page. Page speed and anything else not listed above was not "
        "measured -- say so rather than estimating it.\n"
        # The review is read on screen and pasted into the proposal's landing
        # section, and neither renders Markdown -- so a heading written as
        # "## CTA Status" arrived as literally that. Asked for here, and
        # enforced by clean_ai_text below, for the reason the Smart 1 Labs
        # exclusion is checked rather than merely requested.
        + hub_spec.FORMATTING_DIRECTIVE)
    try:
        review = hub_spec.clean_ai_text(_openai_response(prompt, 5000))
    except Exception as exc:                            # noqa: BLE001
        logger.exception("Landing-page review failed")
        # The reading survives the model failing: what was found on the page
        # is the checkable half, and it is worth showing on its own.
        return jsonify({"ok": False, "url": url, "observed": observed,
                        "error": "The page was read, but the AI review failed",
                        "detail": str(exc)}), 502
    return jsonify({"ok": True, "review": review, "url": observed.get("url") or url,
                    "observed": observed, "summary": _lp.summary_line(observed)})


@app.post("/api/zipcodes-in-radius")
def api_zipcodes_in_radius():
    """The ZIP Codes a radius touches — per target area, not per campaign.

    Same lookup the IO builder runs. Having it here means the ZIP list is
    attached to the area it belongs to while the proposal is still being
    built, rather than being rebuilt at IO time against whichever origin
    happened to be typed first.
    """
    body = request.get_json(force=True) or {}
    origin = str(body.get("origin") or "").strip()
    radius = str(body.get("radius") or "").strip()
    if not origin or not radius:
        return jsonify({"ok": False, "error": "An origin and a radius are required."}), 400
    prompt = (
        f"Find the complete list of United States ZIP Codes whose geographic polygon is fully or "
        f"partially touched by a {radius}-mile radius centered on {origin}. Include a ZIP Code whenever "
        f"any portion of that ZIP Code area intersects the radius, not only when its centroid is inside. "
        "Use current authoritative geographic sources where possible. Return only five-digit ZIP Codes, "
        "comma-separated, sorted ascending, with no commentary. Be exhaustive and do not intentionally "
        "omit any matching ZIP Code.")
    # `search=True`, and the call falls back without the tool rather than
    # failing: the tool riding on this request is what stopped the button.
    try:
        zips = hub_areas.zip_list(_openai_response(prompt, 12000, search=True))
    except Exception as exc:                            # noqa: BLE001
        logger.exception("ZIP-radius lookup failed")
        return jsonify({"ok": False, "error": "ZIP-radius lookup failed",
                        "detail": str(exc)}), 502
    if not zips:
        # Said in the terms of the question that was asked. "No ZIP Codes were
        # returned" reads as a radius with nothing in it, which is not a thing
        # that happens -- it was always the call, never the geography.
        return jsonify({"ok": False,
                        "error": f"The lookup came back with no ZIP Codes for "
                                 f"{radius} miles around {origin}. Check the "
                                 f"origin is a real city or ZIP Code, or enter "
                                 f"the list by hand."}), 502
    return jsonify({"ok": True, "zipcodes": ", ".join(zips), "count": len(zips),
                    "warning": "AI-assisted ZIP-radius results should be reviewed before "
                               "trafficking — ZIP boundaries and radius intersections change."})


@app.get("/api/quotes/<int:qid>/target-map.png")
def quote_target_map(qid):
    """The campaign's coverage map, as the proposal carries it.

    A URL rather than a data URL in the page, because the preview redraws on
    every keystroke of an edited section and a re-encoded PNG in the DOM each
    time is the same picture paid for a hundred times. `hub.target_map`
    caches the composed image, so the PDF built a second later is free.

    `v` is ignored deliberately: it carries the areas' signature so a changed
    campaign gets a new URL and the browser stops serving the old picture.
    A stale map on a proposal is the worst version of this bug -- it is
    plausible, it is dated, and it is about somewhere else.
    """
    db = SessionLocal()
    try:
        q = db.get(Quote, qid)
        if not q:
            return jsonify({"ok": False, "error": "Quote not found"}), 404
        number, raw = q.quote_number, q.data
    finally:
        db.close()
    # Only a malformed blob is caught. A broader except here read an
    # AttributeError on the wrong column name as "quote could not be read",
    # which is a 404 that looks exactly like an empty campaign -- the whole
    # feature silently absent with nothing anywhere saying why.
    try:
        state = json.loads(raw or "{}")
    except (TypeError, ValueError):
        return jsonify({"ok": False, "error": "This quote's data could not be read"}), 404
    png, meta = campaign_map(state)
    if not png:
        # 404 rather than a placeholder image: the builder asks for the reason
        # separately and the client document simply leaves the map out. An
        # "unavailable" graphic is the one outcome that could reach a client.
        return jsonify({"ok": False, "error": meta.get("reason") or "No map"}), 404
    resp = send_file(BytesIO(png), mimetype="image/png",
                     download_name=f"target-map-{number or qid}.png")
    resp.headers["Cache-Control"] = "private, max-age=600"
    return resp


@app.post("/api/target-map/status")
def api_target_map_status():
    """What the map would show, for the builder — never for the client.

    Answers the three questions separately, because they need three different
    things done about them: which areas are drawn, which are covered but not
    drawable (a DMA, a state), and which could not be placed at all and are
    therefore somebody's to fix.
    """
    body = request.get_json(force=True) or {}
    areas = hub_areas.normalize(body.get("areas") or body.get("targetAreas")) \
        or hub_areas.from_legacy(body)
    ok, why = hub_map.available()
    if not ok:
        return jsonify({"ok": True, "measured": False, "reason": why,
                        "plotted": [], "not_plotted": [], "unfound": []})
    try:
        placed = hub_map.locate(areas)
    except Exception as exc:                            # noqa: BLE001
        return jsonify({"ok": True, "measured": False, "plotted": [],
                        "not_plotted": [], "unfound": [],
                        "reason": f"We could not look the areas up ({exc})."}), 200
    drawn = [row["kind"] in hub_map.NOT_DRAWN for row in placed["not_plotted"]]
    return jsonify({
        "ok": True,
        "measured": bool(placed["points"]),
        "plotted": [{"label": p["label"], "found": p["found"],
                     "radius": p["radius"]} for p in placed["points"]],
        # Covered by the campaign and not drawable, versus not found at all.
        # Only the second is a defect, and a screen that lists them together
        # asks somebody to fix a DMA.
        "not_plotted": [r for r, is_kind in zip(placed["not_plotted"], drawn) if is_kind],
        "unfound": [r for r, is_kind in zip(placed["not_plotted"], drawn) if not is_kind],
        # Through hub_map's own reader rather than a sentence written here.
        # This said "an area needs a city or a ZIP Code" whenever nothing was
        # plotted, whichever of the four ways it went -- so it printed directly
        # above the warnbox naming the city the area already carries, and told
        # somebody to fix a DMA that is correctly not drawn.
        "reason": "" if placed["points"] else hub_map.nothing_plotted_reason(placed),
        "attribution": hub_config.map_tile_attribution})


@app.post("/api/paste-areas")
def api_paste_areas():
    """A pasted block of locations, read into target areas.

    Nothing is added here. The rows come back with a sentence per line saying
    how each was read, and the rep presses Add -- because a paste that
    silently assumed a ten-mile radius on eight of twelve lines is eight
    decisions nobody made, and a line this could not read must be seen rather
    than counted.
    """
    body = request.get_json(force=True) or {}
    result = hub_areas.parse_paste(body.get("text"), existing=body.get("areas"))
    return jsonify({"ok": True, **result})


@app.post("/api/paste-places")
def api_paste_places():
    """The same paste, for competitors, venues and places to target."""
    body = request.get_json(force=True) or {}
    result = hub_areas.parse_places(body.get("text"),
                                    kind=str(body.get("kind") or "competitor"),
                                    existing=body.get("existing"))
    return jsonify({"ok": True, **result})


# The research prompt. Separate from the route so the rules are readable and
# so the test harness can assert they are still in it -- three of them are the
# difference between a list a rep can use and a list that puts a business on a
# client's proposal for being plausible.
FIND_TARGETS_RULES = (
    "Rules you must follow:\n"
    "1. Only name businesses, venues or places that genuinely exist today and "
    "are inside or adjacent to the target areas given. Do not invent a name, "
    "and do not list a national brand that has no location there.\n"
    "2. Give a street address ONLY where you are confident it is that "
    "location's real address. If you are not sure, leave the address empty. An "
    "empty address is correct and useful — the campaign can still target the "
    "brand and its customers' behavior. A wrong address builds a geo-fence "
    "around somebody else's building.\n"
    "3. Say in one short line what targeting them is FOR — whose customers "
    "they are and why this client wants them. Not a description of the "
    "business.\n"
    "4. Prefer places the client's actual customers overlap with. A stadium, a "
    "trade show, a hospital campus or a university is worth more than a "
    "competitor nobody visits.\n"
    "5. Return at most 12 rows, strongest first, and no commentary.\n")


@app.post("/api/find-targets")
def api_find_targets():
    """Research competitors, venues and places worth targeting.

    Everything this returns is a **suggestion**. Nothing is added to the
    campaign and nothing reaches the client document until a person ticks it:
    printing a researched list on a proposal is us telling a client who their
    competitors are on a model's say-so, and that is the paragraph a client
    checks hardest. Same rule `modules/ads_builder` applies to its competitor
    research.

    An address is carried only where the model gave one, is labelled
    unverified, and is never derived from the name -- the rule
    `modules/ads_builder/logo.py` works to. "We could not research this" and
    "there is nobody worth naming" are answered differently, because only the
    second means stop looking.
    """
    body = request.get_json(force=True) or {}
    client = str(body.get("client") or "").strip()
    areas = hub_areas.normalize(body.get("areas") or body.get("targetAreas")) \
        or hub_areas.from_legacy(body)
    where = hub_areas.summary(areas, limit=6) or str(body.get("geo") or "").strip()
    if not where:
        return jsonify({"ok": False, "error": "Add at least one target area first — "
                                              "without somewhere to look, this would "
                                              "return national brands."}), 400
    kinds = [k for k in (body.get("kinds") or ["competitor", "venue", "place"])
             if k in hub_areas.PLACE_KINDS] or list(hub_areas.PLACE_KINDS)
    # The two searches the wizard offers are different questions, and the
    # prompt says which is being asked -- the same statement the button shows
    # the rep before they press it, so what the model is told and what the
    # person was promised cannot drift apart.
    focus = ""
    if kinds == ["competitor"]:
        focus = ("Focus: COMPETITORS ONLY. Name companies relevant to this "
                 "client's industry and services -- businesses competing for "
                 "the same customers -- inside the target areas given. Do not "
                 "list venues or general places of interest.\n")
    elif "competitor" not in kinds:
        focus = ("Focus: PLACES ONLY. Name locations whose visitors would be "
                 "a good target audience for this client to reach -- venues, "
                 "campuses, employers, retail parks and event sites where "
                 "their customers already gather -- inside the target areas "
                 "given. Do not list competitor businesses.\n")
    already = [str((r or {}).get("name") or "").strip()
               for r in (body.get("existing") or []) if isinstance(r, dict)]
    prompt = (
        "You are a media planner at Smart 1 Marketing building the targeting "
        "for a local advertising campaign. Research, using the web, who and "
        "where this campaign should go after.\n\n"
        f"Client: {client or 'the advertiser'}\n"
        f"Their website: {str(body.get('url') or '') or 'not given'}\n"
        f"Industry: {str(body.get('industry') or '') or 'not given'}\n"
        f"What they sell: {str(body.get('sells') or '') or 'not given'}\n"
        f"Target areas: {where}\n"
        f"Campaign goals: {', '.join(str(o) for o in (body.get('objectives') or [])) or 'not given'}\n"
        + (f"Already named on this campaign, do not repeat: {', '.join(already[:40])}\n"
           if already else "")
        + "\nReturn STRICT JSON only: {\"targets\":[{\"kind\":\"competitor|venue|place\","
        "\"name\":\"\",\"address\":\"\",\"why\":\"\",\"confidence\":\"confirmed|likely\"}]}\n"
        f"Kinds to include: {', '.join(kinds)}. "
        "competitor = a business that takes this client's customers. "
        "venue = somewhere their customers gather (a stadium, an arena, an "
        "expo center, a campus). place = any other location worth fencing (a "
        "retail park, a hospital, an employer, an event site).\n"
        + focus + "\n" + FIND_TARGETS_RULES)
    try:
        raw = _json_from_ai(_openai_response(prompt, 6000))
    except Exception as exc:                            # noqa: BLE001
        # Deliberately not an empty list: "we could not look" and "there is
        # nobody" send a rep to two different places, and only the second one
        # means the campaign has no conquesting to do.
        return jsonify({"ok": False, "error": "The research request failed",
                        "detail": str(exc)}), 502

    seen = {n.lower() for n in already if n}
    out = []
    for row in (raw.get("targets") or [])[:12]:
        if not isinstance(row, dict):
            continue
        name = str(row.get("name") or "").strip()[:200]
        if not name or name.lower() in seen:
            continue
        seen.add(name.lower())
        kind = str(row.get("kind") or "competitor").strip().lower()
        out.append({
            "kind": kind if kind in hub_areas.PLACE_KINDS else "competitor",
            "name": name,
            "address": str(row.get("address") or "").strip()[:200],
            "note": hub_spec.plain_text(str(row.get("why") or ""))[:200],
            "confidence": ("confirmed" if str(row.get("confidence") or "").lower()
                           == "confirmed" else "likely"),
            # Ticked by a person before it is anything at all.
            "accepted": False,
        })
    return jsonify({
        "ok": True, "targets": out, "searched": where,
        "note": ("Nothing came back for this campaign. That is an answer — it "
                 "does not mean the search failed." if not out else
                 "Researched, not verified. Tick what belongs on this campaign; "
                 "check any address before it is used to build a geo-fence."),
    })


@app.post("/api/zip-rule")
def api_zip_rule():
    """Read a ZIP exception written in words, and apply it to a ZIP list.

    Deliberately a round trip rather than a JavaScript copy of the rule.
    Target areas and the creative classifier each carry a mirror already, and
    each needs its own test proving the two halves still agree; a third --
    carrying a table of every state's ZIP prefixes -- is a mirror that would
    drift silently and be wrong about which state a campaign runs in. The
    browser stores what comes back on the area, so every later read is local
    and `hub.target_areas` stays the only thing that knows what
    "only New Jersey" means.
    """
    body = request.get_json(force=True) or {}
    rule = hub_areas.parse_zip_rule(body.get("text"))
    result = hub_areas.apply_zip_rule(body.get("zips"), rule)
    return jsonify({"ok": True, "rule": rule, "kept": result["kept"],
                    "dropped": result["dropped"], "applied": result["applied"],
                    "note": result["note"],
                    "describe": hub_areas.describe_zip_rule(rule)})


@app.post("/api/estimate-audience")
def api_estimate_audience():
    """Size every target area on the campaign, and total them.

    Each area is sized on its own. A campaign covering Carmel, Fishers and an
    Indianapolis DMA buy is three different reach questions, and one merged
    answer hides that two of them overlap heavily while the third does not.

    Anything the AI cannot size falls back to the shared geometric estimate,
    and an area neither can size is returned as null — "not measured", never
    a confident zero.
    """
    body = request.get_json(force=True) or {}
    areas = hub_areas.normalize(body.get("areas") or body.get("targetAreas")) \
        or hub_areas.from_legacy(body)
    if not areas:
        return jsonify({"ok": False, "error": "Add a target area first."}), 400

    demographics = {"gender": body.get("gender") or "Both",
                    "ages": body.get("ages") or [],
                    "income": body.get("income") or [],
                    "industry": body.get("industry") or ""}
    prompt = (
        "You are a media planner sizing advertising reach for United States geographies. For EACH "
        "target area below, estimate population, addressable audience after the stated demographic "
        "filters, households, and reachable devices. Return STRICT JSON only: "
        "{\"areas\": [{\"id\": \"...\", \"population\": n, \"addressable_audience\": n, "
        "\"households\": n, \"devices\": n, \"rationale\": \"one short sentence\"}]}. "
        "Use whole numbers. If an area cannot be sized from real data, omit it from the array rather "
        "than guessing — an omitted area is reported as not measured, which is correct, and an "
        "invented one is not.\n\n"
        + json.dumps({"demographics": demographics,
                      "areas": [{"id": a["id"], "area": hub_areas.label(a),
                                 "type": a["type"], "radius_miles": a["radius"],
                                 "zip_codes": a["zips"]} for a in areas]},
                     ensure_ascii=False))
    by_id, ai_used = {}, False
    try:
        for row in (_json_from_ai(_openai_response(prompt, 4000)).get("areas") or []):
            if row.get("id"):
                by_id[str(row["id"])] = row
        ai_used = bool(by_id)
    except Exception as exc:                            # noqa: BLE001
        logger.warning("audience estimate fell back to the built-in model: %s", exc)

    out, totals = [], {"pop": 0, "aud": 0, "hh": 0, "dev": 0}
    measured = 0
    # Each area keeps its own full figures -- that is the per-area question
    # this route exists to answer -- and the TOTALS remove the overlap on
    # shared ZIP Codes, because three radii over one metro count the same
    # households three times and the total is the number the reach section
    # prints on the client's document. hub_areas.overlap_factors() is the one
    # reading of how much is shared; an area with no ZIP list counts whole,
    # and the note below says which of the two happened.
    factors = hub_areas.overlap_factors(areas)
    deduped = False
    for area, factor in zip(areas, factors):
        row = by_id.get(area["id"]) or {}
        population = int(row.get("population") or 0) or (hub_areas.estimated_population(area) or 0)
        if not population:
            out.append({"id": area["id"], "label": hub_areas.label(area),
                        "measured": False, "note": "Not measured — this area could "
                                                   "not be sized from what was entered."})
            continue
        audience = int(row.get("addressable_audience") or 0) or int(population * 0.85)
        households = int(row.get("households") or 0) or int(audience / 1.9)
        devices = int(row.get("devices") or 0) or int(audience * 2.3)
        measured += 1
        out.append({"id": area["id"], "label": hub_areas.label(area), "measured": True,
                    "pop": population, "aud": audience, "hh": households, "dev": devices,
                    "ai": bool(row), "rationale": str(row.get("rationale") or "")})
        if factor < 1.0:
            deduped = True
        totals["pop"] += round(population * factor)
        totals["aud"] += round(audience * factor)
        totals["hh"] += round(households * factor)
        totals["dev"] += round(devices * factor)

    note = ""
    if deduped:
        note = ("Where areas share ZIP Codes, the shared ZIPs are counted once "
                "in these totals — the per-area figures above still show each "
                "area whole.")
    elif len(areas) > 1 and any(f == 1.0 and not hub_areas.zip_list(a.get("zips"))
                                for a, f in zip(areas, factors)):
        note = ("An area with no ZIP list cannot be checked for overlap, so "
                "these totals may double-count people that nearby areas share. "
                "Finding the ZIP Codes on each area fixes that.")
    unmeasured = [row["label"] for row in out if not row["measured"]]
    return jsonify({"ok": True, "areas": out, "ai": ai_used,
                    "totals": totals if measured else None,
                    "measured": measured, "unmeasured": unmeasured, "note": note})


# =====================================================================
# Delivery — the PDF the client gets, filed and pushed to Smart 1 Suite
# =====================================================================
def _suite_status():
    try:
        from hub import suite_opportunity
        return suite_opportunity.status()
    except Exception as exc:                            # noqa: BLE001
        return {"ok": False, "problems": [f"Suite helper unavailable ({type(exc).__name__})."]}


@app.get("/api/suite/status")
def api_suite_status():
    """Whether an opportunity will actually be created, in words."""
    return jsonify(_suite_status())


def _upload_proposal_pdf(quote_number, client, pdf_bytes, revision):
    """Put the client-facing PDF somewhere it can be linked to.

    Through hub.storage rather than a local cloudinary.config() call, per the
    migration rule. Returns "" when Cloudinary is not configured — the archive
    copy in the database is still there, and the caller says so rather than
    handing anyone a dead link.
    """
    try:
        from hub import storage
        from hub.config import settings
        safe = storage.slug(f"{quote_number}-{client or 'client'}", "quote")
        public_id = f"{settings.folder('proposals')}/quotes/{safe}-r{revision}.pdf"
        return storage.put("proposals", public_id, pdf_bytes,
                           public_id=public_id, overwrite=True).url or ""
    except Exception as exc:                            # noqa: BLE001
        logger.warning("proposal PDF upload failed: %s", exc)
        return ""


# =====================================================================
# The client's copy of the proposal
#
# Until now a proposal reached a client as a PDF and stopped there: nothing
# knew whether it had been opened, and the status was a pill a rep clicked
# from memory. `/sales/builder/p/<token>` is the document a client opens, and
# the one thing they can do on it is accept it.
#
# Deliberately **no edits from the client**. Smart 1 Ads offers three answers
# because a campaign estimate is a thing to negotiate line by line; a proposal
# is a document somebody says yes to, and a change request arriving here would
# be a second inbox for something the rep is already having a conversation
# about. A client who wants something different says so to the rep, who edits
# the quote and sends the same link again.
#
# The page **embeds the PDF** rather than re-rendering the proposal in HTML.
# That is the whole reason it is cheap: the PDF, the Word export and the
# preview are already three renderers of one document, and a fourth would be
# the drift this codebase pays for twice already. The client reads exactly the
# document that was signed off, and there is nothing to keep in step.
# =====================================================================
# wsgi.py hands this to BOTH the AuthGuard (so a client with no Hub login can
# reach it) and HubBar (so the sidebar, help layer and feedback tab are not
# injected into a document a client reads). One list, so the mount and the
# module cannot disagree about what is public -- the arrangement
# modules/scans and modules/ads_builder already use.
PUBLIC_PREFIXES = ("/p/", "/api/p/")

_TOKEN_BYTES = 16       # 128 bits: the token is the whole security model


def _client_base() -> str:
    """The origin a client-facing link is built from.

    `PUBLIC_BASE_URL` first, because a link is pasted into an email and has to
    work from outside; `request.host_url` only as the fallback. Trimmed to an
    origin for the reason `modules/image_picker/provisioning.py` gives: a
    dispatcher-mounted module's url_root carries its own mount, and pasting a
    path onto it builds /sales/builder/sales/builder/p/… -- a 404 the client
    meets and nobody else does.
    """
    base = (hub_config.public_base_url or "").strip()
    if not base:
        try:
            base = request.host_url
        except Exception:                       # noqa: BLE001
            base = ""
    base = (base or "").rstrip("/")
    # PUBLIC_BASE_URL has held a whole callback URL before now, so trim to the
    # origin rather than trusting it.
    match = re.match(r"^(https?://[^/]+)", base)
    return match.group(1) if match else base


def share_url(token: str) -> str:
    return f"{_client_base()}/sales/builder/p/{token}"


def _get_or_make_share(db, q, sending: bool = False) -> "QuoteShare":
    """This quote's client link, minted once and kept.

    `sending` marks the revision the rep deliberately sent, which is what lets
    the panel say "you have edited this since you sent it" rather than leaving
    a rep to remember.
    """
    share = (db.query(QuoteShare)
             .filter(QuoteShare.quote_id == q.id)
             .order_by(QuoteShare.id.desc()).first())
    if share is None:
        share = QuoteShare(quote_id=q.id, token=secrets.token_urlsafe(_TOKEN_BYTES),
                           created_by=_signed_in_as(), sent_revision=q.revision or 1)
        db.add(share)
        db.flush()
    if sending:
        share.revoked_at = None
        share.sent_revision = q.revision or 1
        share.sent_at = datetime.now(timezone.utc)
    return share


def _staff_reader() -> str:
    """The signed-in Hub user, or "" for a client.

    The rule the whole feature was asked for with: a rep has to be able to
    open the client's link and read the document without marking it read. A
    session cookie is what tells the two apart, and it costs the rep nothing
    to remember.
    """
    try:
        from hub import auth as hub_auth
        return hub_auth.verify_cookie_value(
            request.cookies.get(hub_auth.COOKIE_NAME)) or ""
    except Exception:                           # noqa: BLE001
        return ""


def _share_state(db, q) -> dict:
    """Everything the builder's panel says about the client's copy."""
    share = (db.query(QuoteShare)
             .filter(QuoteShare.quote_id == q.id)
             .order_by(QuoteShare.id.desc()).first())
    if share is None:
        return {"shared": False, "url": "", "views": 0, "opens": [],
                "accepted": None, "superseded": None, "revision": q.revision or 1}
    views = (db.query(QuoteView)
             .filter(QuoteView.quote_id == q.id)
             .order_by(QuoteView.at.desc()).limit(50).all())
    accepts = (db.query(QuoteAcceptance)
               .filter(QuoteAcceptance.quote_id == q.id)
               .order_by(QuoteAcceptance.id.desc()).all())
    current = q.revision or 1
    live = next((a for a in accepts if (a.revision or 1) == current), None)
    superseded = next((a for a in accepts if (a.revision or 1) != current), None)

    def _acc(row):
        if not row:
            return None
        return {"name": row.name, "email": row.email,
                "revision": row.revision or 1,
                "at": row.at.isoformat() if row.at else ""}

    return {
        "shared": True,
        "revoked": bool(share.revoked_at),
        "url": share_url(share.token),
        "token": share.token,
        "revision": current,
        "sent_revision": share.sent_revision or 1,
        # "You have edited it since you sent it" is the thing a rep cannot
        # know from a list of opens, and the reason to re-send the link.
        "edited_since_sent": bool(share.sent_at) and (share.sent_revision or 1) != current,
        "sent_at": share.sent_at.isoformat() if share.sent_at else "",
        "views": len(views),
        "views_this_revision": sum(1 for v in views if (v.revision or 1) == current),
        "first_view": views[-1].at.isoformat() if views else "",
        "last_view": views[0].at.isoformat() if views else "",
        "opens": [{"at": v.at.isoformat() if v.at else "",
                   "revision": v.revision or 1, "device": v.device or ""}
                  for v in views[:12]],
        "accepted": _acc(live),
        "superseded": _acc(superseded) if not live else None,
        "validity": _validity_block(_quote_window(q, _state_of(q), share.sent_at)),
    }


@app.get("/api/quotes/<int:qid>/share")
def api_share_state(qid):
    db = SessionLocal()
    try:
        q = db.get(Quote, qid)
        if not q:
            return jsonify({"ok": False, "error": "Quote not found"}), 404
        return jsonify({"ok": True, "share": _share_state(db, q)})
    finally:
        db.close()


@app.post("/api/quotes/<int:qid>/share")
def api_share_create(qid):
    """Mint (or re-send) the client link for this quote."""
    db = SessionLocal()
    try:
        q = db.get(Quote, qid)
        if not q:
            return jsonify({"ok": False, "error": "Quote not found"}), 404
        share = _get_or_make_share(db, q, sending=True)
        if q.status == "Draft":
            q.status = "Sent"
        log_activity(db, q.id, "🔗", f"Client link sent — revision {q.revision or 1}")
        db.commit()
        # `actor` is the wrapper's own keyword -- audit.for_module passes it
        # from its actor_fn -- so who did this travels under its own name. The
        # first-positional trap this file's own docstring names, one keyword on.
        _audit("quote_shared", client=q.client, quote=q.quote_number,
               sent_by=_signed_in_as(), revision=q.revision or 1)
        return jsonify({"ok": True, "share": _share_state(db, q)})
    finally:
        db.close()


@app.post("/api/quotes/<int:qid>/validity")
def api_share_validity(qid):
    """How long this quote's pricing stands, if not the house window.

    Written into the quote's own data blob -- never a new column, the
    `create_all()` rule -- and set from the share panel because that is where
    the send happens and where the answer matters. Clearing it (0 or blank)
    puts the quote back on the house window rather than removing the window,
    which is a different thing and not one this button offers.
    """
    body = request.get_json(silent=True) or {}
    raw = body.get("days")
    db = SessionLocal()
    try:
        q = db.get(Quote, qid)
        if not q:
            return jsonify({"ok": False, "error": "Quote not found"}), 404
        state = _state_of(q)
        if raw in (None, "", 0, "0"):
            state.pop("validityDays", None)
            chosen = 0
        else:
            try:
                chosen = int(raw)
            except (TypeError, ValueError):
                return jsonify({"ok": False,
                                "error": "Give the number of days as a whole number."}), 400
            if not (hub_validity.MIN_DAYS <= chosen <= hub_validity.MAX_DAYS):
                # Refused by name rather than silently clamped: a rep who
                # typed 3650 and got 365 has been told something different
                # from what they asked for, on a date a client relies on.
                return jsonify({"ok": False, "error":
                                f"A proposal can stand for between "
                                f"{hub_validity.MIN_DAYS} and "
                                f"{hub_validity.MAX_DAYS} days."}), 400
            state["validityDays"] = chosen
        q.data = json.dumps(state, ensure_ascii=False)
        log_activity(db, q.id, "⏰",
                     (f"Pricing held for {chosen} days" if chosen
                      else "Pricing back on the standard window"))
        db.commit()
        return jsonify({"ok": True, "share": _share_state(db, q)})
    finally:
        db.close()


@app.post("/api/quotes/<int:qid>/share/revoke")
def api_share_revoke(qid):
    db = SessionLocal()
    try:
        q = db.get(Quote, qid)
        if not q:
            return jsonify({"ok": False, "error": "Quote not found"}), 404
        share = (db.query(QuoteShare).filter(QuoteShare.quote_id == q.id)
                 .order_by(QuoteShare.id.desc()).first())
        if share:
            share.revoked_at = datetime.now(timezone.utc)
            log_activity(db, q.id, "🔒", "Client link revoked")
            db.commit()
        return jsonify({"ok": True, "share": _share_state(db, q)})
    finally:
        db.close()


# ------------------------------------------------------------------ public
def _open_share(db, token):
    """The quote behind a token, or (None, None).

    Revoked, deleted and never-existed all answer the same 404, because a
    client-facing URL that says "this one expired" tells somebody probing
    which tokens are real -- the rule `modules/ads_builder` settled.
    """
    share = db.query(QuoteShare).filter(QuoteShare.token == str(token or "")).first()
    if not share or share.revoked_at:
        return None, None
    q = db.get(Quote, share.quote_id)
    if not q:
        return None, None
    return share, q


@app.get("/p/<token>")
def page_client_proposal(token):
    db = SessionLocal()
    try:
        share, q = _open_share(db, token)
        if not share:
            return render_template("client_gone.html"), 404
        staff = _staff_reader()
        state = _state_of(q)
        win = _quote_window(q, state, share.sent_at)
        return render_template(
            "client_proposal.html", q=q, token=token, staff=staff,
            revision=q.revision or 1,
            # Past the date the page says so and names who to ask rather than
            # 404ing. A revoked or invented token answers 404 because saying
            # "that one expired" tells somebody probing which tokens are real;
            # an expired quote is a real quote belonging to a real client who
            # is trying to say yes.
            validity=win, validity_note=hub_validity.client_note(win),
            expired=bool(win.get("expired")),
            contact_name=str(state.get("salesContact") or q.salesperson or ""),
            contact_email=str(state.get("salesEmail") or ""),
            accepted=_share_state(db, q)["accepted"])
    finally:
        db.close()


@app.get("/p/<token>.pdf")
def page_client_pdf(token):
    """The document itself, for the embed and for the download button.

    Deliberately no view recorded here. The page reports its own open once,
    and counting this too would report every client as having read it twice --
    and would count the rep's own check of the link, which is the one thing
    the feature was asked not to do.
    """
    db = SessionLocal()
    try:
        share, q = _open_share(db, token)
        if not share:
            return jsonify({"ok": False, "error": "Not found"}), 404
        state = json.loads(q.data or "{}")
        ensure_sections(state)
        # The share row is right here, so the validity line is dated from the
        # send rather than from a second lookup that could disagree with the
        # page the client just read it on.
        pdf_bytes, title = build_proposal_pdf(q, state, sent_at=share.sent_at)
    except Exception as exc:                    # noqa: BLE001
        logger.warning("client PDF failed for %s: %s", token, exc)
        return jsonify({"ok": False, "error": "Not found"}), 404
    finally:
        db.close()
    resp = send_file(BytesIO(pdf_bytes), mimetype="application/pdf",
                     as_attachment=request.args.get("download") == "1",
                     download_name=title + ".pdf")
    resp.headers["Cache-Control"] = "private, max-age=60"
    return resp


@app.post("/api/p/<token>/opened")
def api_client_opened(token):
    """The page reporting that a browser actually rendered it.

    Called from the page rather than recorded on the HTML request, because a
    mail security gateway fetches every link in a message within seconds of
    delivery and runs no JavaScript. Counted there, every proposal reads as
    opened the moment it is sent -- a confident wrong answer that stops a rep
    chasing a client who has never seen it.

    Answers 200 whatever it decides. This is a beacon: a client's page must
    not show an error because we chose not to count their visit.
    """
    db = SessionLocal()
    try:
        share, q = _open_share(db, token)
        if not share:
            return jsonify({"ok": True, "counted": False, "reason": "unknown link"})
        staff = _staff_reader()
        if staff:
            return jsonify({"ok": True, "counted": False,
                            "reason": "staff preview — not counted"})
        agent = request.headers.get("User-Agent", "")
        automated, why = hub_views.looks_automated(agent, request.headers)
        if automated:
            return jsonify({"ok": True, "counted": False, "reason": why})
        visitor = hub_views.visitor_hash(
            _client_ip(), hub_config.secret_key)
        # The window is per visitor **per revision**. A reload of the same
        # document is one read; the first sight of a revision the rep has just
        # sent is a new one, whatever the clock says -- and "have they opened
        # the version I sent this morning" is the whole question a revised
        # quote asks.
        last = (db.query(QuoteView)
                .filter(QuoteView.quote_id == q.id,
                        QuoteView.visitor == visitor,
                        QuoteView.revision == (q.revision or 1))
                .order_by(QuoteView.at.desc()).first())
        last_seen = last.at.replace(tzinfo=timezone.utc).timestamp() if last and last.at else None
        now = datetime.now(timezone.utc)
        if not hub_views.counts_as_new_view(last_seen, now.timestamp()):
            return jsonify({"ok": True, "counted": False,
                            "reason": "already counted this visit"})
        db.add(QuoteView(quote_id=q.id, token=share.token, at=now,
                         revision=q.revision or 1, visitor=visitor,
                         device=hub_views.device_kind(agent)))
        log_activity(db, q.id, "👁", f"Client opened revision {q.revision or 1}")
        db.commit()
        return jsonify({"ok": True, "counted": True})
    except Exception as exc:                    # noqa: BLE001
        logger.warning("view beacon failed for %s: %s", token, exc)
        return jsonify({"ok": True, "counted": False, "reason": "not recorded"})
    finally:
        db.close()


@app.post("/api/p/<token>/accept")
def api_client_accept(token):
    """The client says yes, to one specific revision.

    Name and email are required and are the whole of it: an acceptance nobody
    can attribute is not an acceptance, and anything more is a form standing
    between a client and the word yes. Stamped with the revision, because
    approving is a statement about a specific document -- if the quote is
    revised afterwards the panel says the acceptance was superseded rather
    than quietly carrying it forward onto a document nobody agreed to.
    """
    body = request.get_json(silent=True) or {}
    name = str(body.get("name") or "").strip()[:200]
    email = str(body.get("email") or "").strip()[:200]
    if not name or "@" not in email:
        return jsonify({"ok": False,
                        "error": "Please give your name and an email address so "
                                 "we know who accepted."}), 400
    db = SessionLocal()
    try:
        share, q = _open_share(db, token)
        if not share:
            return jsonify({"ok": False, "error": "This link is no longer active."}), 404
        if _staff_reader():
            # A rep pressing Accept on the client's page would file an
            # acceptance in a client's name that the client never gave.
            return jsonify({"ok": False,
                            "error": "You are signed in to the Hub — a proposal is "
                                     "accepted by the client, on their own copy."}), 403
        state = _state_of(q)
        win = _quote_window(q, state, share.sent_at)
        if win.get("expired"):
            # Server-side as well as on the form: a rule the page keeps while
            # the write breaks it is not a rule. 409 rather than 404 -- the
            # link is genuinely theirs and the answer is "not at this price",
            # which is a different thing from "no such quote".
            out = hub_validity.refusal(
                win, str(state.get("salesContact") or q.salesperson or ""),
                str(state.get("salesEmail") or ""))
            _audit("quote_accept_expired", quote=q.quote_number,
                   client=q.client, expired_on=win.get("expires_on", ""))
            return jsonify({"ok": False, **out}), 409
        revision = q.revision or 1
        existing = (db.query(QuoteAcceptance)
                    .filter(QuoteAcceptance.quote_id == q.id,
                            QuoteAcceptance.revision == revision).first())
        if existing:
            return jsonify({"ok": True, "already": True,
                            "accepted": {"name": existing.name,
                                         "at": existing.at.isoformat() if existing.at else ""}})
        db.add(QuoteAcceptance(
            quote_id=q.id, token=share.token, revision=revision,
            name=name, email=email,
            visitor=hub_views.visitor_hash(_client_ip(), hub_config.secret_key)))
        q.status = "Approved"
        log_activity(db, q.id, "✅", f"Accepted by {name} — revision {revision}")
        db.commit()
        _audit("quote_accepted", client=q.client, quote=q.quote_number,
               accepted_by=f"{name} <{email}>", revision=revision)
        return jsonify({"ok": True, "accepted": {"name": name, "revision": revision}})
    finally:
        db.close()


def _client_ip() -> str:
    try:
        from hub import auth as hub_auth
        return hub_auth.client_ip(request.headers, request.remote_addr or "")
    except Exception:                           # noqa: BLE001
        return request.remote_addr or ""


@app.post("/api/quotes/<int:qid>/deliver")
def deliver_quote(qid):
    """Send a proposal to a client: PDF, filed on the client, opportunity in Suite.

    Three things happen, in an order chosen so a later failure never loses an
    earlier success:

      1. The branded PDF is generated and archived on the quote (always works).
      2. It is uploaded and filed on the client's Hub record, so it shows up on
         Client 360 with everything else that client has been sent.
      3. An opportunity is created in Smart 1 Suite.

    Step 3 is the one that can come back asking a question. If Suite holds no
    contact for this client, the response is `needs_contact` — the proposal is
    already saved and filed at that point, and the rep supplies a name plus an
    email or phone and posts again. That is the whole reason this is not done
    silently in the background: inventing a contact from a business name is
    how a Suite account fills with records nobody can call.
    """
    body = request.get_json(silent=True) or {}
    db = SessionLocal()
    try:
        q = db.get(Quote, qid)
        if not q:
            return jsonify({"ok": False, "error": "Quote not found"}), 404
        state = json.loads(q.data or "{}")
        ensure_sections(state)

        pdf_bytes, title = build_proposal_pdf(q, state)
        q.pdf_blob = pdf_bytes
        q.pdf_filename = title + ".pdf"
        q.pdf_generated_at = datetime.now(timezone.utc)

        url = q.pdf_url
        if not url or body.get("regenerate", True):
            url = _upload_proposal_pdf(q.quote_number, q.client, pdf_bytes, q.revision or 1) or q.pdf_url
        q.pdf_url = url or ""

        filed, filed_note = None, ""
        if q.client:
            try:
                from hub import proposals as hub_proposals
                # One filed document per revision. Pressing Send twice on the
                # same revision -- a double click, or answering the Suite
                # contact question -- must not leave the client's record
                # showing the same proposal three times; each of those rows
                # looks like a separate thing we sent them.
                prior_id, _, prior_rev = (q.client_filed_as or "").partition("@")
                if prior_id and prior_rev == str(q.revision or 1):
                    hub_proposals.delete_proposal(q.client, prior_id)
                filed = hub_proposals.add_proposal(
                    q.client, title + ".pdf", pdf_bytes,
                    title=f"{q.quote_number} — {q.package or 'Marketing Proposal'}"
                          + (f" (rev {q.revision})" if (q.revision or 1) > 1 else ""),
                    note=f"Built in the Proposal Builder. {q.geo_summary or ''}".strip(),
                    actor=request.environ.get("s1hub.user") or state.get("salesContact") or "",
                    value=str(q.monthly_budget or 0), term="monthly", status="sent")
                q.client_filed_as = f"{filed.get('id') or ''}@{q.revision or 1}"[:64]
            except Exception as exc:                    # noqa: BLE001
                logger.warning("filing the proposal on the client failed: %s", exc)
                filed_note = (f"The proposal could not be filed on {q.client}'s record "
                              f"({type(exc).__name__}). It is still saved here.")
        else:
            filed_note = "No client name on this quote, so it was not filed on a client record."

        contact = body.get("contact") or {
            "name": state.get("clientContactName") or "",
            "email": state.get("clientContactEmail") or "",
            "phone": state.get("clientContactPhone") or "",
        }
        if q.suite_contact_id and not contact.get("id"):
            contact["id"] = q.suite_contact_id

        # The PDF is built and filed by this point. A Suite problem -- even the
        # helper failing to import -- reports itself; it never costs the work
        # already done.
        try:
            from hub import suite_opportunity
            suite = suite_opportunity.push_proposal(
                client=q.client, title=f"{q.client} — {q.package or 'Marketing'} Proposal "
                                       f"({q.quote_number})",
                value=float(q.monthly_budget or 0), contact=contact,
                website=q.website, pdf_url=q.pdf_url,
                opportunity_id=q.suite_opportunity_id or "")
        except Exception as exc:                        # noqa: BLE001
            logger.warning("Suite push failed: %s", exc)
            suite = {"ok": False, "reason": f"Smart 1 Suite is unreachable "
                                            f"({type(exc).__name__})."}
        if suite.get("ok"):
            q.suite_opportunity_id = str(suite.get("opportunity_id") or "")[:64]
            found = suite.get("contact") or {}
            q.suite_contact_id = str(found.get("id") or "")[:64]
            # Remember the contact, so converting this quote to an IO does not
            # ask for it a second time. setdefault would not do: the wizard
            # seeds these keys as empty strings, so they are present-but-blank
            # rather than absent.
            for key, value in (("clientContactName", found.get("name")),
                               ("clientContactEmail", found.get("email")),
                               ("clientContactPhone", found.get("phone"))):
                if value and not state.get(key):
                    state[key] = value
            q.data = json.dumps(state, ensure_ascii=False)

        q.delivered_at = datetime.now(timezone.utc)
        if q.status == "Draft":
            q.status = "Sent"
        log_activity(db, q.id, "📤",
                     f"{q.quote_number} delivered — {q.client}"
                     + (f", Suite opportunity {q.suite_opportunity_id}" if suite.get("ok")
                        else ", no Suite opportunity"))
        db.commit()
        try:
            _audit("proposal_delivered", client=q.client, quote=q.quote_number,
                   suite=bool(suite.get("ok")))
        except Exception:                               # noqa: BLE001
            pass
        return jsonify({"ok": True, "quote": quote_json(q), "pdf_url": q.pdf_url,
                        "filed": filed, "filed_note": filed_note, "suite": suite})
    finally:
        db.close()


# =====================================================================
# The retired standalone builder's archive
# =====================================================================
@app.get("/api/legacy/proposals")
def legacy_proposals():
    """Proposals built in the old /sales/proposals tool.

    Read-only, and listed beside the quotes rather than merged into them: they
    have no quote number, no revision and no campaign data, and giving them
    invented ones would make the pipeline read as bigger and better-specified
    than it is.
    """
    try:
        from modules.proposal_builder import store as legacy_store
    except Exception as exc:                            # noqa: BLE001
        return jsonify({"ok": True, "proposals": [], "available": False,
                        "note": f"The old builder's archive is unreadable ({type(exc).__name__})."})
    try:
        rows = legacy_store.search_proposals(
            q=request.args.get("q", ""),
            limit=clamp_int(request.args.get("limit"), 50, 1, 200))
    except Exception as exc:                            # noqa: BLE001
        logger.warning("legacy proposal index unreadable: %s", exc)
        return jsonify({"ok": True, "proposals": [], "available": False,
                        "note": "The old builder's archive could not be read."})
    return jsonify({"ok": True, "proposals": rows, "available": True})


@app.post("/api/legacy/proposals/<pid>/import")
def legacy_import(pid):
    """Reopen an old standalone proposal as a quote in this builder.

    What carries over is what was structured: the customer, the industry, the
    recommended package and its monthly price, and the narrative blocks as
    editable text sections. What does not carry over is a media plan, because
    the old proposals never had one — the products were named inside prose.
    The imported quote therefore lands as a Draft with its gaps listed, which
    is an honest description of it.
    """
    try:
        from modules.proposal_builder import store as legacy_store
    except Exception:                                   # noqa: BLE001
        return jsonify({"ok": False, "error": "The old builder's archive is unavailable."}), 503
    record = legacy_store.get_proposal(pid)
    if not record:
        return jsonify({"ok": False, "error": "No such proposal."}), 404

    customer = record.get("customer") or {}
    sections, monthly = [], record.get("monthly_investment") or 0
    for i, block in enumerate(record.get("blocks") or []):
        data = block.get("data") or {}
        title = str(data.get("heading") or data.get("title") or "").strip()
        parts = [str(data.get("body") or "").strip()]
        for item in (data.get("items") or []):
            if isinstance(item, dict):
                parts.append(f"• {item.get('label', '')}: {item.get('value', '')}")
            else:
                parts.append(f"• {item}")
        text = "\n".join(p for p in parts if p)
        if title or text:
            sections.append({"id": f"legacy{i}", "title": title or f"Section {i + 1}",
                             "kind": "text", "enabled": True, "body": text})

    city, st = customer.get("city", ""), customer.get("state", "")
    where = ", ".join(x for x in (city, st) if x) or customer.get("zip", "")
    areas = hub_areas.normalize([{"name": where, "origin": where}]) if where else []

    state = {
        "client": customer.get("business_name", ""),
        "url": customer.get("website", ""),
        "industry": record.get("industry_label") or record.get("industry") or "",
        "description": "", "descriptionMode": "I have one",
        "objectives": [], "kpis": [], "exclusions": [], "audiences": [], "items": [],
        "targetAreas": areas,
        "budget": int(float(monthly or 0)), "months": 12,
        "clientContactName": customer.get("contact_name", ""),
        "clientContactEmail": customer.get("contact_email", ""),
        "clientContactPhone": customer.get("contact_phone", ""),
        "salesContact": customer.get("salesperson", ""),
        "sections": sections or None,
        "importedFrom": f"proposal-builder:{pid}",
    }
    hub_areas.apply_to_legacy(state, areas)

    db = SessionLocal()
    try:
        q = Quote(quote_number=next_quote_number(db), status="Draft",
                  data=json.dumps(state, ensure_ascii=False))
        summarize_into(q, state)
        q.package = str(record.get("recommended_package") or "")[:40]
        q.pdf_url = str(record.get("pdf_url") or "")[:600]
        db.add(q)
        db.flush()
        log_activity(db, q.id, "📥",
                     f"{q.quote_number} imported from the old Proposal Builder — {q.client}")
        db.commit()
        return jsonify({"ok": True, "quote": quote_json(q, include_data=True),
                        "note": "Goals, products and budget were never structured in the old "
                                "builder, so they are blank rather than guessed. The copy came "
                                "across as editable sections."})
    finally:
        db.close()


@app.errorhandler(Exception)
def handle_unexpected_error(exc):
    logger.exception("Unhandled application error")
    return jsonify({"error": "Internal server error", "type": type(exc).__name__, "message": str(exc)}), 500


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.getenv("PORT", "8001")), debug=False)
